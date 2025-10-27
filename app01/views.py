import base64
from datetime import timedelta, datetime, time
from django.shortcuts import render
from rest_framework.response import Response
from app01.utils.workTime import clean_and_calculate_attendance
from rest_framework.decorators import api_view
from rest_framework.views import APIView
from app01 import models
import uuid #用于生成token
from app01.serializers.repair import companyInfoSerializer,RepairOrderSerializer,\
    RepairAdviceSerializer,ReportSerializer,machineInfoSerializer,workerInfoSerializer,QuotationSerializer
from django.core.files.storage import default_storage
import os
from django.http import FileResponse
from djangoProject import settings
import requests
from django.core.files.storage import FileSystemStorage
import io

def getOpendID(code):
    url = f'https://api.weixin.qq.com/sns/jscode2session?appid={settings.WX_APPID}&secret={settings.WX_APPSECRET}&js_code={code}'
    response = requests.get(url)
    resData = response.json()
    openid = resData.get('openid')
    return openid


#用户登录
class loginView(APIView):
    authentication_classes = []
    def post(self,request):
        print(request.data)
        data = request.data
        role = data.get("role")
        workerName = None
        # print(role)
        # 验证用户身份
        if role == 1: #工作人员
            userName = data.get("userName")
            pwd = data.get("password")
            # print(userName)
            #检查数据库
            log_object = models.workerInfo.objects.filter(Tele=userName).first() #搜索账号
            log_pwd = models.workerInfo.objects.filter(password=pwd).first()
            workerName = log_pwd.name
            print(workerName)
            print([log_object,log_pwd])
            if not log_object or not log_pwd:
                return Response({'status':False,'message':"账号或密码错误"})
        else:
            companyName = data.get("companyName")
            code = data.get("code")
            # 检查数据库
            log_object = models.companyInfo.objects.filter(name=companyName).first()  # 搜索账号
            openid = log_object.openid
            if not log_object:
                return Response({'status': False, 'message': "该企业暂未注册，请前往注册"})
            elif not openid:
                # url = f'https://api.weixin.qq.com/sns/jscode2session?appid={settings.WX_APPID}&secret={settings.WX_APPSECRET}&js_code={code}'
                # response = requests.get(url)
                # resData = response.json()
                openid = getOpendID(code)
                log_object.openid = openid

        token = str(uuid.uuid4())
        log_object.token = token
        log_object.save()
        data = {}
        data['token'] = token
        if workerName:
            data['workerName'] = workerName
        return Response({'status': True, 'message': "success",'data':data,'Tele': log_object.Tele})



#用户注册
class registerView(APIView):
    authentication_classes = []
    def get(self,request):
        return Response({'status':True,'message':"success"})
    def post(self,request):
        print(request.data)
        data = request.data
        role = data.get("role")
        # 验证用户身份
        if role == 'worker': #工作人员
            userName = data.get("userName")
            userTele = data.get("userTele")
            pwd = data.get("password")
            code = data.get('code')
            signCode = data.get('signCode')
            if signCode == "qazwsx121219":
                # print([userName,pwd,role])
                #存入数据库
                sign_object = models.workerInfo.objects.filter(name=userName).first()
                if sign_object:
                    return Response({'status':False,'message':"用户名已注册"})
                else:
                    worker = models.workerInfo.objects.create(name=userName,Tele=userTele,password=pwd)
                    # url = f'https://api.weixin.qq.com/sns/jscode2session?appid={settings.WX_APPID}&secret={settings.WX_APPSECRET}&js_code={code}'
                    # response = requests.get(url)
                    # resData = response.json()
                    # openid = resData.get('openid')
                    worker.openid = getOpendID(code)
                    worker.save()

                return Response({'status':True,'message':"success"})
            else:
                return Response({'status':False,'message':"注册码错误，无法注册"})
        else: #客户
            companyName = data.get("companyName")
            contactInfo = data.get("contactInfo")
            code = data.get('code')
            print([companyName,contactInfo])
            #存入数据库
            sign_object = models.companyInfo.objects.filter(name=companyName).first()
            if sign_object:
                return Response({'status':False,'message':"该企业已注册"})
            else:
                company = models.companyInfo.objects.create(name=companyName,Tele=contactInfo)
                # url = f'https://api.weixin.qq.com/sns/jscode2session?appid={settings.WX_APPID}&secret={settings.WX_APPSECRET}&js_code={code}'
                # response = requests.get(url)
                # resData = response.json()
                # openid = resData.get('openid')
                company.openid = getOpendID(code)
                company.save()
            return Response({'status':True,'message':"success"})

#下拉框辅助登录
class searchCompany(APIView):
    authentication_classes = []
    def get(self,request):
        data = request.query_params
        print(data)
        keyword = data.get("keyword")
        companies = models.companyInfo.objects.filter(name__icontains=keyword).all()
        results = companyInfoSerializer(instance=companies,many=True)

        return Response(results.data)

#报修时的机器选项
class getMachine(APIView):
    def post(self,request):
        data = request.data
        print(data)
        companyName = data.get("companyName")
        company = models.companyInfo.objects.filter(name=companyName).first()
        machineObjects = company.machine.all()
        machines = {'value':'', 'options':[]}
        machinesId = {'value':'', 'options':[]}
        for i in machineObjects:
            machines['options'].append({'value':i.name,'label':i.name})
            machinesId['options'].append({'value':i.machine_id,'label':i.machine_id})
        # print(machinesId['options'])
        if machines['options'] == []:
            return Response({'status':False,'message':"您暂未购买机器，请联系工作人员"})
        return Response({'status':True,'message':"success",'machines':machines,'machinesId':machinesId})



from django.utils import timezone
#创建报修单
class RepairOrder(APIView):
    def get(self,request):
        return Response({'status':True,'message':"success"})
    def post(self, request):
        # 处理表单数据
        data = request.data
        # 报修单号
        time_part = timezone.now().strftime("%Y%m%d%H%M")
        random_part = uuid.uuid4().hex[:4].upper()  # 取UUID前4位并大写
        order_id = f"{time_part}_{random_part}"
        data["order_id"]= order_id
        companyId = models.companyInfo.objects.filter(name=data.get("created_by")).first().id
        valid_fields = {f.name for f in models.RepairOrder._meta.get_fields()}  # 获取模型所有字段名
        filtered_data = {k: v for k, v in data.items() if k in valid_fields and v != ''}
        filtered_data['company_id'] = companyId
        print(filtered_data)
        new_order = models.RepairOrder.objects.create(**filtered_data)
        if new_order:
            return Response({'status': True})
        else:
            return Response({'status': False})
#上传报修图片
class RepairImage(APIView):
    def post(self,request):
        print(request.data)
        uploaded_file = request.data.get("image")
        print(uploaded_file)
        method = request.data.get("method")
        if method == "delete":
            # 删除图片
            file_url = request.data.get("filePath")
            file_path = file_url.replace(settings.url_ROOT+'/', '')  # 去掉前缀
            if file_path:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    return Response({'status': True, 'message': 'Image deleted successfully.'})
                except Exception as e:
                    return Response({'status': False, 'message': str(e)})
            else:
                return Response({'status': False, 'message': 'File path not provided.'})
        if uploaded_file:
            name = request.data.get("name")
            companyId = models.companyInfo.objects.filter(name=name).first().id
            # 存储图片到本地/服务器
            date_path = timezone.now().strftime("repair_images/%Y/%m/%d")
            # 生成唯一文件名（防止冲突）
            file_ext = os.path.splitext(uploaded_file.name)[1]  # 文件扩展名
            new_filename = f"{uuid.uuid4()}{file_ext}"  # 随机文件名

            file_path = default_storage.save(f'{date_path}/{companyId}/{new_filename}', uploaded_file)
            file_url = default_storage.url(file_path)
            file_url = settings.url_ROOT + file_url
            print(file_url)
            return Response({'status': True, 'url': file_url})
        else:
            return Response({'status': True})

class showQuotation(APIView):
    def post(self,request):
        data = request.data
        orderID = data.get('orderID')
        repairOrder = models.RepairOrder.objects.get(order_id = orderID)
        quotation = repairOrder.Quotations.first()
        data = QuotationSerializer(instance=quotation)
        # print(data)
        # print(data['quotation'])
        print(type(data['quotation']))
        data = data['quotation'].value
        return Response({'status':True,'data':data})
#上传建议图片
class AdviceImage(APIView):
    # authentication_classes = []
    def post(self,request):
        print(request.data)
        uploaded_file = request.data.get("image")
        name = request.data.get("name")
        role = request.data.get("role")
        print(role)
        if role == '0':
            Id = models.companyInfo.objects.filter(name=name).first().id
        else:
            Id = models.workerInfo.objects.filter(name=name).first().id
        #存储图片到本地/服务器
        date_path = timezone.now().strftime("advice_images/%Y/%m/%d")
        # 生成唯一文件名（防止冲突）
        file_ext = os.path.splitext(uploaded_file.name)[1]  # 文件扩展名
        new_filename = f"{uuid.uuid4().hex[:6]}{file_ext}"  # 随机文件名
        if role == 0:
            file_path = default_storage.save(f'{date_path}/customer/{Id}/{new_filename}', uploaded_file)
        else:
            file_path = default_storage.save(f'{date_path}/worker/{Id}/{new_filename}', uploaded_file)
        file_url = default_storage.url(file_path)
        file_url = settings.url_ROOT + file_url
        # print(file_url)
        return Response({'status':True,'url':file_url})
#创建建议表单
class RepairAdvice(APIView):
    def post(self,request):
        data = request.data
        repairOrder = data.get('repair_order_id')
        if repairOrder:
            id = models.RepairOrder.objects.filter(order_id=repairOrder).first().id
            data['repair_order_id'] = id
        # valid_fields = {f.name for f in models.RepairOrder._meta.get_fields()}  # 获取模型所有字段名
        # filtered_data = {k: v for k, v in data.items() if k in valid_fields and v != ''}
        new_order = models.RepairAdvice.objects.create(**data)
        if new_order:
            return Response({'status':True})
        else:
            return Response({'status':False})

class registerPwd(APIView):
    def post(self,request):
        data = request.data
        companyName=data.get("companyName")
        pwd = data.get("sign_password")
        user = models.companyInfo.objects.filter(name = companyName).first()
        otherUser = models.companyInfo.objects.filter(password = pwd).first()
        if user.password:
            return Response({'status':False,'message':'密码已注册，找回密码请联系管理员'})
        elif otherUser:
            return Response({'status':False,'message':'改密码已被注册，请重新设置'})
        user.password = pwd
        user.save()
        return Response({'status':True})

class loginHistory(APIView):
    def post(self,request):
        companyName = request.data.get("companyName")
        pwd = request.data.get("history_password")
        code = request.data.get("code")
        openid = getOpendID(code)
        user = models.companyInfo.objects.filter(name=companyName,password=pwd).first()

        if user:
            user.openid = openid
            user.save()
            return Response({'status':True})
        else:
            return Response({'statue':False,'message':'密码错误'})

class repairHistory(APIView):
    def post(self,request):
        data = request.data
        # print(data)
        name = data.get("companyName")
        status = data.get("status")
        quotationStatus = data.get("quotationStatus")
        print(quotationStatus)
        # print(name)
        companyId = models.companyInfo.objects.filter(name=name).first().id
        if status == 0:
            orders = models.RepairOrder.objects.filter(company_id=companyId,status=status,quotation_status=quotationStatus).all()
        else:
            orders = models.RepairOrder.objects.filter(company_id=companyId,status=status).all()
        orderInfo = RepairOrderSerializer(instance=orders, many=True)
        orderdata = orderInfo.data

        for i in range(len(orderdata)):
            id = models.RepairOrder.objects.filter(order_id=orderdata[i]['order_id']).first().id

            worker = models.RepairOrder.objects.get(order_id=orderdata[i]['order_id']).worker
            advice = models.RepairAdvice.objects.filter(repair_order_id=id).first()
            if advice:
                orderdata[i]['advice'] = True
            else:
                orderdata[i]['advice'] = False
            if worker:
                name = worker.name
                Tele = worker.Tele
                orderdata[i]['worker'] = name
                orderdata[i]['workerWeichat'] = Tele


            if status == 3:
                reports = models.RepairOrder.objects.get(company_id=companyId,status=status,id = id).reports.all()
                reportInfo = ReportSerializer(instance=reports, many=True)
                print(reportInfo.data)
                del orderdata[i]['repair_images']
                orderdata[i]['finishedFiles'] = reportInfo.data[0]['finished_images']
                orderdata[i]['report'] = reportInfo.data[0]['report']
                # orderdata[i]['worker_id'] = models.
            print(orderdata)

        return Response({'status':True,'data':orderdata})

class adviceHistory(APIView):
    def post(self,request):
        data = request.data
        print(data)
        name = data.get("Name")
        status = data.get("status")
        role = data.get("")
        if status == 0:
            orders=models.RepairAdvice.objects.filter(created_by=name).all()
        else:
            orders = models.RepairAdvice.objects.filter(created_by=name,status=status).all()
        orderInfo = RepairAdviceSerializer(instance=orders,many=True)
        print(orderInfo.data)
        # print(orderInfo.data)

        return Response({'status':True,'data':orderInfo.data})


class downloadFile(APIView):
    def get(self,request):
        path = settings.MEDIA_ROOT + '/quotation'
        token = request.query_params.get('token')
        print(token)
        order_id = request.GET.get('id')
        orderID = models.RepairOrder.objects.filter(order_id=order_id).first().id
        file_path = path + f'/{orderID}/{order_id}.pdf'  # 文件绝对路径
        file_name = os.path.basename(file_path)
        response = FileResponse(open(file_path, 'rb'))
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response

class workerRepair(APIView):
    def post(self,request):
        data = request.data
        status = int(data.get("status"))
        print(status)
        workerPhone = int(data.get("workerId"))
        workerID = models.workerInfo.objects.filter(Tele=workerPhone).first()
        RepairOrders = models.RepairOrder.objects.filter(worker_id=workerID,status=status).all()
        RepairInfo = RepairOrderSerializer(instance=RepairOrders,many=True)
        print(RepairInfo.data)
        return Response({'status':True,'data':RepairInfo.data})

class workerChangePassword(APIView):
    def post(self,request):
        data = request.data
        workerPhone = int(data.get("workerId"))
        pwd = data.get("password")
        worker = models.workerInfo.objects.filter(Tele=workerPhone).first()
        if worker:
            worker.password = pwd
            worker.save()
            return Response({'status':True})
        else:
            return Response({'status':False,'message':'工人不存在'})

class workerRepairDetail(APIView):
    def post(self,request):
        data = request.data
        orderId = data.get("orderId")
        order = models.RepairOrder.objects.filter(order_id=orderId).first()
        RepairInfo = RepairOrderSerializer(instance=order)
        data = RepairInfo.data
        company = order.company
        adress = company.adress
        data['adress'] = adress
        return Response({'status':True,'data':data})

class changeRepairStatus(APIView):
    def post(self,request):
        data = request.data
        status=int(data.get("status"))
        orderId = data.get("orderId")
        print(orderId)
        print(status)
        order = models.RepairOrder.objects.filter(order_id=orderId).first()
        order.status = status
        process_date = timezone.now().strftime('%Y-%m-%d')
        print(process_date)
        order.process_date = process_date
        order.save()

        # 发送完成通知
        openid = order.company.openid
        order_id = order.order_id
        status = "处理中"
        tag = "尊敬的客户，您的报修订单我们已开始处理"
        template_id = "v9k5A7NCfPgHiinemRO2yAw_LzXgmHrnud7XVrOpE2M"  # 进度通知模板
        page = f"/packageA/pages/repair_history/repair_history"
        data = {
            "character_string1": {"value": order_id},
            "phrase3": {"value": status},  # 根据实际模板调整
            "thing5": {"value": tag}
        }

        WeiChatAPI.send_subscribe_message(openid, template_id, page, data)

        return Response({'status':True})

class reportImages(APIView):
    def post(self,request):
        # print(request.data)
        uploaded_file = request.data.get("image")
        if uploaded_file:
            name = request.data.get("name")
            # 存储图片到本地/服务器
            date_path = timezone.now().strftime("report_images/%Y/%m/%d")
            # 生成唯一文件名（防止冲突）
            file_ext = os.path.splitext(uploaded_file.name)[1]  # 文件扩展名
            new_filename = f"{uuid.uuid4().hex[:6]}{file_ext}"  # 随机文件名

            file_path = default_storage.save(f'{date_path}/{name}/{new_filename}', uploaded_file)
            file_url = default_storage.url(file_path)
            file_url = settings.url_ROOT + file_url
            # print(file_url)
            return Response({'status': True, 'url': file_url})
        else:
            return Response({'status': True})

class repairReport(APIView):
    def post(self,request):
        data = request.data
        print(data)
        orderId = data.get('repairId')
        repairOrder = models.RepairOrder.objects.filter(order_id=orderId).first()
        repairId = repairOrder.id
        repairOrder.status = 3
        finished_date = timezone.now().strftime('%Y-%m-%d')
        repairOrder.finished_date = finished_date
        repairOrder.save()
        report = data.get('ReportDetail')
        images = data.get('images')
        models.Report.objects.create(repairOrder_id=repairId,report=report,finished_images=images)

        #发送完成通知
        openid = repairOrder.company.openid
        order_id = repairOrder.order_id
        status = "已完成"
        tag = "尊敬的客户，您的报修订单已完成"
        template_id = "v9k5A7NCfPgHiinemRO2yAw_LzXgmHrnud7XVrOpE2M"  # 进度通知模板
        page = f"/packageA/pages/repair_history/repair_history"
        data = {
            "character_string1": {"value": order_id},
            "phrase3": {"value": status},  # 根据实际模板调整
            "thing5": {"value": tag}
        }

        WeiChatAPI.send_subscribe_message(openid, template_id, page, data)
        return Response({'status':True})

#临时文件存储
class tempImage(APIView):
    def post(self,request):

        method = request.data.get("method")
        if method == 'delete':
            # 删除临时文件
            file_url = request.data.get("filePath")
            file_path = file_url.replace(settings.url_ROOT+'/', '')  # 去掉URL前缀
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                return Response({'status': True, 'message': '文件删除成功'})
            except Exception as e:
                return Response({'status': False, 'message': '文件删除失败', 'error': str(e)})
        uploaded_file = request.data.get("tempImage")
        if uploaded_file:
            # 存储图片到本地/服务器
            date_path = timezone.now().strftime("temp_images/%Y/%m/%d")
            # 生成唯一文件名（防止冲突）
            file_ext = os.path.splitext(uploaded_file.name)[1]  # 文件扩展名
            new_filename = f"{uuid.uuid4()}{file_ext}"  # 随机文件名

            file_path = default_storage.save(f'{date_path}/{new_filename}', uploaded_file)
            file_url = default_storage.url(file_path)
            file_url = settings.url_ROOT + file_url
            return Response({'status': True, 'url': file_url})
        else:
            return Response({'status': False, 'message': '文件上传失败'})


#前后端不分离 后台管理系统页面
from django.shortcuts import render,redirect
from django import forms
from app01.utils.pagination import Pagination
from app01.BootstrapForm.repair import c_info,c_info_reform,\
    machine_info,worker_info,repair_info,report_info,advice_info,reply_info
from model_creat import generate_quotation
import pandas as pd
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.db import transaction
from app01.utils.weichat_api import WeiChatAPI

def authentication(token):
    manager = models.ManagerInfo.objects.filter(token=token).first()
    return manager
from django.db.models import Q

# 常见视频文件扩展名
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.avi', '.mkv', '.flv', '.wmv', '.webm'}

def is_video(url):
    return any(url.lower().endswith(ext) for ext in VIDEO_EXTENSIONS)

# 页面设计
class Login(forms.Form):
    account = forms.CharField(
        label="用户账号",
        widget=forms.TextInput(attrs = {"class":"form-control"}),
        required= True
    )
    password = forms.CharField(
        label="密码",
        widget=forms.TextInput(attrs={"class": "form-control", "type": "password"}),
        required= True
    )
    identity = forms.IntegerField(
        label = "身份",
        widget = forms.Select(choices=((0,"用户"),(1,"管理者")) , attrs= {"class":"form-control"}),
    )


class managerLogin(APIView):
    authentication_classes = []
    def get(self,request):
        form = Login()
        error = request.query_params.get('error')
        print(request.data)
        return render(request,'common/login.html',{"form":form,"error":error})
    def post(self,request):
        data = request.data
        # print(data)
        form = Login(data = request.data)
        account = data.get("account")
        if account == '':
            form.add_error("account","请输入账号")
            return render(request,'common/login.html',{"form":form})
        password = data.get("password")
        if password == '':
            print(password)
            form.add_error("password","请输入账号")
            return render(request,'common/login.html',{"form":form})
        token = data.get("csrfmiddlewaretoken")
        object = models.ManagerInfo.objects.filter(account = account,password=password).first()
        if object:
            object.token = token
            object.save()
            # print(object)
            return redirect(f'/c_info_list/?token={token}')
        else:
            form.add_error("password","账号或密码错误，请重试")
            return render(request,'common/login.html',{"form":form})

class companyInfoView(APIView):
    # authentication_classes = []
    def get(self,request):
        token = request.query_params.get('token')
        manager = authentication(token)
        page = request.query_params.get('page')
        params = request.META.get('QUERY_STRING', '')
        if not manager:
            # form.add_error("password","认证失败，请重新登录")
            error = '认证失败，请重新登录'
            return redirect(f'/?error={error}')
        search_dict = {}
        queryset = models.companyInfo.objects.all()
        # 3. 处理关键词模糊搜索
        if search := request.GET.get('search'):

            # 创建Q对象组合多个字段的模糊搜索
            q_objects = Q()
            search_fields = [
                'name__icontains',
                'manager__icontains',
            ]

            for field in search_fields:
                q_objects |= Q(**{field: search})

            # 将Q对象应用到queryset
            queryset = queryset.filter(q_objects)


        # 4. 应用其他筛选条件
        if search_dict:
            queryset = queryset.filter(**search_dict)

        # 5. 排序（可选）
        queryset = queryset.order_by('id')
        # queryset = models.companyInfo.objects.filter(**search_dict).order_by('id')
        # print(queryset)
        page_object = Pagination(request, queryset)
        # print(page_object.page_queryset)
        # query = companyInfoSerializer(instance=queryset,many=True)
        queryset = page_object.page_queryset
        query = companyInfoSerializer(instance=queryset, many=True)
        # print(query.data)
        if not page:
            page = 1
        context = {
            # 'queryset': page_object.page_queryset,
            'queryset': query.data,
            'page_string': page_object.html(),
            'token':token,
            # 'page':page,
            'params': params,
        }
        return render(request, 'userInfo/c_info_list.html', context)
    def post(self,request):
        # token = request.query_params.get('token')
        # page = request.data.get('page')
        params = request.META.get('QUERY_STRING', '')
        # print(page)
        return redirect(f'/c_info_list/?{{params}}')



class companyInfoReform(APIView):
    # authentication_classes = []
    def get(self,request,id):
        token = request.query_params.get('token')
        company = models.companyInfo.objects.filter(id = id).first()
        data = companyInfoSerializer(instance=company)
        # page = request.query_params.get('page')
        # print(data.data['machine'])
        machine = data.data['machine']
        form = c_info_reform(instance=company)
        # print(form)
        machineObject = models.machine.objects.all()
        machines = machineInfoSerializer(instance=machineObject,many=True)
        # print(machines.data)
        params = request.META.get('QUERY_STRING', '')
        context = {
            'form':form,
            'machine':machine,
            'all_machines':machines.data,
            "token":token,
            "companyId":id,
            # 'page':page
            'params': params,
        }
        return render(request,'userInfo/c_info_reform.html',context)
    def post(self,request,id):
        data = request.data
        # print(data)
        # token = request.query_params.get('token')
        machines = data.get('machines')
        # page = request.query_params.get('page')
        if machines:
            company = models.companyInfo.objects.get(id = id)
            for item in machines:
                machine = models.machine.objects.filter(machine_id=item['machine_id'])
                print(machine)
                if machine:
                    machine = models.machine.objects.get(machine_id=item['machine_id'])
                    company.machine.add(machine)
                else:
                    print("创建新的machine")
                    machine_object = models.machine.objects.create(name=item['name'], machine_id=item['machine_id'])
                    machine = models.machine.objects.get(id=machine_object.id)
                    company.machine.add(machine)
            return Response({'success': True})
        object = models.companyInfo.objects.filter(id = id)
        valid_fields = {f.name for f in models.companyInfo._meta.get_fields()}  # 获取模型所有字段名
        filtered_data = {k: v for k, v in data.items() if k in valid_fields}
        print(filtered_data)
        object.update(**filtered_data)
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/c_info_list/?{params}')
class companyMachineImport(APIView):
    def post(self,request):
        token = request.query_params.get('token')
        id = request.query_params.get('id')
        company = models.companyInfo.objects.filter(id = id).first()
        if 'excel_file' not in request.FILES:
            return Response({'success': False, 'message': '请上传文件'}, status=400)
        try:
            excel_file = request.FILES['excel_file']

            # 读取Excel文件
            df = pd.read_excel(excel_file)

            # 验证必要列是否存在
            required_columns = ['机器型号', '机器编号']
            for col in required_columns:
                if col not in df.columns:
                    return Response({'success': False, 'message': f'缺少必要列: {col}'}, status=400)
            imported_count = 0
            errors = []

            with transaction.atomic():
                for index, row in df.iterrows():
                    # 创建企业信息
                    # object = companyInfo.objects.filter(name=name[0])
                    # if not object:
                    try:
                        machineName= row.get('机器型号', '')
                        machineID = row.get('机器编号', '')

                        # 处理所购机器（多对多关系）
                        machineObject = models.machine.objects.filter(name=machineName,machine_id=machineID).first()
                        if not machineObject:
                            machineObject = models.machine.objects.create(name=machineName,machine_id=machineID)
                        if not company.machine.filter(id=machineObject.id).exists():
                            company.machine.add(machineObject)
                        imported_count += 1
                    except Exception as e:
                        errors.append(f"第{index + 2}行错误: {str(e)}")
                        continue

            if errors:
                return Response({
                    'success': True,
                    'imported_count': imported_count,
                    'message': f'成功导入{imported_count}条，但有部分错误:\n' + '\n'.join(errors[:5])  # 只返回前5个错误
                })

            return Response({
                'success': True,
                'imported_count': imported_count,
                'message': f'成功导入{imported_count}条企业信息'
            })

        except Exception as e:
            return Response({
                'success': False,
                'message': f'处理文件时出错: {str(e)}'
            }, status=500)

class companyInfoAdd(APIView):
    def get(self,request):
        form = c_info()
        token = request.query_params.get('token')
        machineObject = models.machine.objects.all()
        machines = machineInfoSerializer(instance=machineObject, many=True)
        params = request.META.get('QUERY_STRING', '')
        return render(request,'userInfo/c_info_add.html',{'form':form,'all_machines':machines.data,'params':params,'token':token})
    def post(self,request):
        data = request.data
        print(data)
        machines = data.get('machines')
        token = request.query_params.get('token')
        print(machines)

        if machines:
            models.companyInfo.objects.create(name='temple_companyInfo')
            for item in machines:
                company = models.companyInfo.objects.get(name='temple_companyInfo')
                machine = models.machine.objects.filter(machine_id = item['machine_id'])
                print(machine)
                if machine:
                    machine = models.machine.objects.get(machine_id = item['machine_id'])
                    company.machine.add(machine)
                else:
                    print("创建新的machine")
                    machine_object = models.machine.objects.create(name = item['name'],machine_id=item['machine_id'])
                    machine = models.machine.objects.get(id = machine_object.id)
                    company.machine.add(machine)
            return Response({'success': True})

        object = models.companyInfo.objects.filter(name='temple_companyInfo')
        valid_fields = {f.name for f in models.companyInfo._meta.get_fields()}  # 获取模型所有字段名
        filtered_data = {k: v for k, v in data.items() if k in valid_fields}
        params = request.META.get('QUERY_STRING', '')
        if object:
            print(filtered_data)
            name = filtered_data['name']
            form = c_info(request.POST)
            print(name)
            if name == '':
                form.add_error('name',"请填写企业名称")
                token = request.query_params.get('token')
                machineObject = models.machine.objects.all()
                machines = machineInfoSerializer(instance=machineObject, many=True)
                object.delete()
                return render(request,'userInfo/c_info_add.html',{'form':form,'all_machines':machines.data,'params':params,'token':token})
            object.update(**filtered_data)
        else:
            name = filtered_data['name']
            form = c_info(request.POST)
            print(name)
            if name == '':
                form.add_error("name", "请填写企业名称")
                token = request.query_params.get('token')
                machineObject = models.machine.objects.all()
                machines = machineInfoSerializer(instance=machineObject, many=True)
                return render(request, 'userInfo/c_info_add.html',
                              {'form': form, 'all_machines': machines.data, 'params':params,'token': token})
            models.companyInfo.objects.create(**filtered_data)
        # return Response({'success':True})
        return redirect(f'/c_info_add/?{params}')

class companyInfoImport(APIView):
    def post(self,request):
        from .models import companyInfo, machine
        if 'excel_file' not in request.FILES:
            return Response({'success': False, 'message': '请上传文件'}, status=400)
        try:
            excel_file = request.FILES['excel_file']

            # 读取Excel文件
            df = pd.read_excel(excel_file)

            # 验证必要列是否存在
            required_columns = ['企业名称', '联系电话']
            for col in required_columns:
                if col not in df.columns:
                    return Response({'success': False, 'message': f'缺少必要列: {col}'}, status=400)
            imported_count = 0
            errors = []

            with transaction.atomic():
                for index, row in df.iterrows():
                    name=row['企业名称']
                    # 创建企业信息
                    object = companyInfo.objects.filter(name=name)
                    if not object:
                        try:

                            company = companyInfo.objects.create(
                                name=row['企业名称'],
                                Tele=row['联系电话'],
                                email=row['联系邮箱'],
                                manager=row['管理者姓名'],
                                gender=1 if str(row.get('性别', '男')).strip() == '男' else 2,
                                adress=row['企业地址'],
                                password=row['密码'],
                            )
                            machineName= row.get('机器型号', '')
                            machineID = row.get('机器编号', '')
                            # if machineName == 'nan':
                            #     continue
                            if pd.isna(machineName) or not str(machineName).strip():
                                continue
                            # 处理所购机器（多对多关系）
                            machineObject = machine.objects.filter(name=machineName,machine_id=machineID).first()
                            if not machineObject:
                                machineObject = machine.objects.create(name=machineName,machine_id=machineID)
                            machine = machine.objects.get(id = machineObject.id)
                            company.machine.add(machineObject)

                            imported_count += 1
                        except Exception as e:
                            errors.append(f"第{index + 2}行错误: {str(e)}")
                            continue
                    else:
                        object.update(
                            name=row['企业名称'],
                            Tele=row['联系电话'],
                            email=row.get('联系邮箱', ''),
                            manager=row.get('管理者姓名', ''),
                            gender=1 if str(row.get('性别', '男')).strip() == '男' else 2,
                            adress=row.get('企业地址', ''),
                            password=row.get('密码', ''),
                        )
                        machineName = row.get('机器型号', '')
                        machineID = row.get('机器编号', '')
                        # if machineName == 'nan':
                        #     continue
                        if pd.isna(machineName) or not str(machineName).strip():
                            continue
                        company = companyInfo.objects.get(id=object.first().id)  # 重新查询
                        # 处理所购机器（多对多关系）
                        machineObject = machine.objects.filter(name=machineName, machine_id=machineID).first()
                        if not machineObject:
                            machineObject = machine.objects.create(name=machineName, machine_id=machineID)

                        # company = object.first()
                        if not company.machine.filter(id=machineObject.id).exists():
                            company.machine.add(machineObject)
                            company.save()
                        imported_count += 1

            if errors:
                return Response({
                    'success': True,
                    'imported_count': imported_count,
                    'message': f'成功导入{imported_count}条，但有部分错误:\n' + '\n'.join(errors[:5])  # 只返回前5个错误
                })

            return Response({
                'success': True,
                'imported_count': imported_count,
                'message': f'成功导入{imported_count}条企业信息'
            })

        except Exception as e:
            print(e)
            return Response({
                'success': False,
                'message': f'处理文件时出错: {str(e)}'
            }, status=500)
class companyInfoDelete(APIView):
    # authentication_classes = []
    def get(self,request,id):
        print(id)
        object = models.companyInfo.objects.filter(id = id)
        token = request.query_params.get('token')
        page = request.query_params.get('page')
        object.delete()
        return redirect(f"/c_info_list/?token={token}&page={page}")

class MachineDelete(APIView):
    def post(self,request):
        data = request.data
        id = data.get('customer_id')
        machine_id  = data.get('machine_id')
        company = models.companyInfo.objects.get(id = id)
        machine = models.machine.objects.get(machine_id=machine_id)
        company.machine.remove(machine)
        return Response({'success':True})


class machineInfoView(APIView):
    # authentication_classes = []
    def get(self,request):
        token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        manager = authentication(token)
        if not manager:
            # form.add_error("password","认证失败，请重新登录")
            error = '认证失败，请重新登录'
            return redirect(f'/?error={error}')
        search_dict = {}
        queryset = models.machine.objects.all()
        # 3. 处理关键词模糊搜索
        if search := request.GET.get('search'):

            # 创建Q对象组合多个字段的模糊搜索
            q_objects = Q()
            search_fields = [
                'machine_id__icontains',
                'name__icontains'
            ]

            for field in search_fields:
                q_objects |= Q(**{field: search})

            # 将Q对象应用到queryset
            queryset = queryset.filter(q_objects)


        # 4. 应用其他筛选条件
        if search_dict:
            queryset = queryset.filter(**search_dict)

        # 5. 排序（可选）
        queryset = queryset.order_by('id')
        # queryset = models.machine.objects.filter(**search_dict).order_by('id')
        print(queryset)
        page_object = Pagination(request, queryset)
        # print(page_object.page_queryset)
        queryset = page_object.page_queryset
        query = machineInfoSerializer(instance=queryset, many=True)
        print(query.data)
        context = {
            # 'queryset': page_object.page_queryset,
            'queryset': query.data,
            'page_string': page_object.html(),
            'token':token,
            'params': params,
        }

        return render(request, 'machineInfo/machine_info.html', context)
    def post(self,request):
        # token = request.query_params.get('token')
        # page = request.data.get('page')
        # print(page)
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/machine_info/?{params}')

class machineInfoDelete(APIView):
    def get(self,request,id):
        object = models.machine.objects.filter(id = id).first()
        # token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        object.delete()
        return redirect(f"/machine_info/?{params}")

class machineInfoAdd(APIView):
    def get(self,request):
        form = machine_info()
        token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return render(request,'machineInfo/machine_info_add.html',{'form':form,'params':params,'token':token})
    def post(self,request):
        data = request.data
        print(data)
        # token = request.query_params.get('token')
        form = machine_info(request.POST)
        name = data.get('name')
        params = request.META.get('Query_String', '')
        machine_id = data.get('machine_id')
        if name == '':
            form.add_error('name',"请输入机器名称")
            if machine_id == '':
                form.add_error('machine_id',"请输入机器编号")
            return render(request, 'machineInfo/machine_info_add.html', {'form': form,'params': params})
        if machine_id == '':
            form.add_error('machine_id', "请输入机器编号")
            return render(request,'machineInfo/machine_info_add.html',{'form':form,'params': params})
        models.machine.objects.create(name = name,machine_id = machine_id)
        return redirect(f'/machine_info_add/?{params}')
class machineInfoImport(APIView):
    def post(self, request):
        from .models import companyInfo, machine
        if 'excel_file' not in request.FILES:
            return Response({'success': False, 'message': '请上传文件'}, status=400)
        try:
            excel_file = request.FILES['excel_file']

            # 读取Excel文件
            df = pd.read_excel(excel_file)

            # 验证必要列是否存在
            required_columns = ['机器型号', '机器编号']
            for col in required_columns:
                if col not in df.columns:
                    return Response({'success': False, 'message': f'缺少必要列: {col}'}, status=400)
            imported_count = 0
            errors = []

            with transaction.atomic():
                for index, row in df.iterrows():
                    machineName = row.get('机器型号', '')
                    machineID = row.get('机器编号', '')
                    object = machine.objects.filter(name=machineName,machine_id=machineID)
                    if not object:
                        try:
                            machine.objects.create(name=machineName, machine_id=machineID)

                            imported_count += 1
                        except Exception as e:
                            errors.append(f"第{index + 2}行错误: {str(e)}")
                            continue

            if errors:
                return Response({
                    'success': True,
                    'imported_count': imported_count,
                    'message': f'成功导入{imported_count}条，但有部分错误:\n' + '\n'.join(errors[:5])  # 只返回前5个错误
                })

            return Response({
                'success': True,
                'imported_count': imported_count,
                'message': f'成功导入{imported_count}条企业信息'
            })

        except Exception as e:
            return Response({
                'success': False,
                'message': f'处理文件时出错: {str(e)}'
            }, status=500)

class workerInfoView(APIView):

    def get(self, request):
        token = request.query_params.get('token')
        manager = authentication(token)
        page = request.query_params.get('page')
        params = request.META.get('QUERY_STRING', '')
        if not manager:
            # form.add_error("password","认证失败，请重新登录")
            error = '认证失败，请重新登录'
            return redirect(f'/?error={error}')
        search_dict = {}
        queryset = models.workerInfo.objects.all()
        # 3. 处理关键词模糊搜索
        if search := request.GET.get('search'):

            # 创建Q对象组合多个字段的模糊搜索
            q_objects = Q()
            search_fields = [
                'name__icontains',
            ]

            for field in search_fields:
                q_objects |= Q(**{field: search})

            # 将Q对象应用到queryset
            queryset = queryset.filter(q_objects)

        # 4. 应用其他筛选条件
        if search_dict:
            queryset = queryset.filter(**search_dict)

        # 5. 排序（可选）
        queryset = queryset.order_by('id')
        # queryset = models.workerInfo.objects.filter(**search_dict).order_by('id')
        # print(queryset)
        page_object = Pagination(request, queryset)
        # print(page_object.page_queryset)
        # query = companyInfoSerializer(instance=queryset,many=True)
        queryset = page_object.page_queryset
        query = workerInfoSerializer(instance=queryset, many=True)
        # print(query.data)
        if not page:
            page = 1
        context = {
            # 'queryset': page_object.page_queryset,
            'queryset': query.data,
            'page_string': page_object.html(),
            # 'page':page,
            'token':token,
            'params': params,
        }
        return render(request, 'workerInfo/worker_list.html', context)
    def post(self,request):
        # token = request.query_params.get('token')
        # page = request.data.get('page')
        # print(page)
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/worker_info/?{params}')

class workerInfoAdd(APIView):
    def get(self,request):
        form = worker_info()
        token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return render(request,'workerInfo/worker_info_add.html',{'form':form,'token':token,'params':params})
    def post(self,request):
        data = request.data
        print(data)
        name = data.get('name')
        Tele = data.get('Tele')
        password = data.get('password')
        models.workerInfo.objects.create(name = name,Tele=Tele,password = password)
        # token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/worker_info_add/?{params}')
class workerInfoReform(APIView):
    def get(self,request,id):
        worker = models.workerInfo.objects.filter(id = id).first()
        workerInfo = worker_info(instance=worker)
        print(workerInfo)
        # page = request.query_params.get('page')
        token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return render(request,'workerInfo/worker_info_reform.html',{'form':workerInfo,'token':token,'workerId':id,'params':params})
    def post(self,request,id):
        data =request.data
        # print(data)
        Tele = data.get('Tele')
        worker = models.workerInfo.objects.filter(Tele=Tele)
        # token = request.query_params.get('token')
        # page = request.query_params.get('page')
        valid_fields = {f.name for f in models.workerInfo._meta.get_fields()}  # 获取模型所有字段名
        filtered_data = {k: v for k, v in data.items() if k in valid_fields}
        print(filtered_data)
        worker.update(**filtered_data)
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/worker_info/{id}/reform/?{params}')

class workerInfoDelete(APIView):
    def get(self,request,id):
        worker = models.workerInfo.objects.filter(id = id)
        worker.delete()
        # page = request.query_params.get('page')
        # token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/worker_info/?{params}')

class workerInfoImport(APIView):
    def post(self, request):
        from .models import workerInfo
        if 'excel_file' not in request.FILES:
            return Response({'success': False, 'message': '请上传文件'}, status=400)
        try:
            excel_file = request.FILES['excel_file']

            # 读取Excel文件
            df = pd.read_excel(excel_file)

            # 验证必要列是否存在
            required_columns = ['姓名', '联系电话','密码']
            for col in required_columns:
                if col not in df.columns:
                    return Response({'success': False, 'message': f'缺少必要列: {col}'}, status=400)
            imported_count = 0
            errors = []

            with transaction.atomic():
                for index, row in df.iterrows():
                    name = row.get('姓名', '')
                    # 创建企业信息
                    object = workerInfo.objects.filter(name=name)
                    if not object:
                        try:
                            workerInfo.objects.create(
                                name=name,
                                Tele=row.get('联系电话',''),
                                password=row.get('密码','')
                                                      )

                            imported_count += 1
                        except Exception as e:
                            errors.append(f"第{index + 2}行错误: {str(e)}")
                            continue

            if errors:
                return Response({
                    'success': True,
                    'imported_count': imported_count,
                    'message': f'成功导入{imported_count}条，但有部分错误:\n' + '\n'.join(errors[:5])  # 只返回前5个错误
                })

            return Response({
                'success': True,
                'imported_count': imported_count,
                'message': f'成功导入{imported_count}条企业信息'
            })

        except Exception as e:
            return Response({
                'success': False,
                'message': f'处理文件时出错: {str(e)}'
            }, status=500)

class repairOrderInfoView(APIView):
    def get(self, request):
        data = request.data
        # print(context['token'])
        token = request.query_params.get('token')
        manager = authentication(token)
        # print(request.query_params)
        if not manager:
            # form.add_error("password","认证失败，请重新登录")
            error = '认证失败，请重新登录'
            return redirect(f'/?error={error}')
        queryset = models.RepairOrder.objects.all()
        search_dict = {}
        # 1. 处理精确匹配的筛选条件
        exact_match_fields = {
            'emergency_level': 'emergency_level',
            'status': 'status',
            'dispatch_status': 'dispatch_status',
            'quotation_status': 'quotation_status'
        }

        for param, field in exact_match_fields.items():
            if value := request.GET.get(param):
                search_dict[field] = value
        # 2. 处理日期范围筛选
        if start_date := request.GET.get('start_date'):
            search_dict['repair_date__gte'] = start_date
        if end_date := request.GET.get('end_date'):
            search_dict['repair_date__lte'] = end_date

        # 3. 处理关键词模糊搜索
        if search := request.GET.get('search'):

            # 创建Q对象组合多个字段的模糊搜索
            q_objects = Q()
            search_fields = [
                'customer_name__icontains',
                'responsible_name__icontains',
                'machine_id__icontains',
                'machine_model__icontains',
                'order_id__icontains',
            ]

            for field in search_fields:
                q_objects |= Q(**{field: search})

            # 将Q对象应用到queryset
            queryset = queryset.filter(q_objects)

            # 搜索外键worker
            worker = models.workerInfo.objects.filter(name=search)
            if worker:
                worker = models.workerInfo.objects.get(name=search)
                queryset = worker.repairOrders.filter(**search_dict)

        # 4. 应用其他筛选条件
        if search_dict:
            queryset = queryset.filter(**search_dict)

        # 5. 排序（可选）
        queryset = queryset.order_by('-repair_date')
        # queryset = models.RepairOrder.objects.filter(**search_dict).order_by('id')
        # print(queryset)
        page_object = Pagination(request, queryset)
        # print(page_object.page_queryset)
        # query = companyInfoSerializer(instance=queryset,many=True)
        queryset = page_object.page_queryset
        query = RepairOrderSerializer(instance=queryset, many=True)
        # print(query.data)
        page = request.query_params.get('page')
        if not page:
            page = 1
        params = request.META.get('QUERY_STRING', '')
        context = {
            # 'queryset': page_object.page_queryset,
            'queryset': query.data,
            'page_string': page_object.html(),
            # 'page':page,
            'token':token,
            'params':params
        }

        return render(request, 'repairOrder/repairOrder_info_list.html', context)
    def post(self,request):
        token = request.query_params.get('token')
        page = request.data.get('page')
        print(page)
        return redirect(f'/repairOrder_info/?token={token}&page={page}')

class repairOrderInfoDetail(APIView):
    def get(self,request,id):
        repairOrder = models.RepairOrder.objects.filter(id = id).first()
        repairInfo = repair_info(instance=repairOrder)
        for field in repairInfo.fields.values():
            field.widget.attrs['readonly'] = True
            field.widget.attrs['disabled'] = True
            field.widget.attrs['style'] = 'background-color: #ffffff;'
        token = request.query_params.get('token')
        all_images =  repairOrder.repair_images
        images = [url for url in all_images if not is_video(url)]
        print(images)
        videos = [url for url in all_images if is_video(url) ]
        print(videos)
        reports = repairOrder.reports.all()
        isReport = False
        reportID = 0
        # page = request.query_params.get('page')
        orderID = repairOrder.id
        if reports:
            isReport = True
            reportID = reports[0].id
        params = request.META.get('QUERY_STRING', '')
        context = {'form':repairInfo,
                   'token':token,
                   'images':images,
                   'videos':videos,
                   'isReport':isReport,
                   "reportID":reportID,
                   'params':params,
                   'orderID': orderID,
                   }
        return render(request,'repairOrder/repair_info_detail.html',context)

class repairOrderInfoDelete(APIView):
    def get(self,request,id):
        repairOrder = models.RepairOrder.objects.filter(id = id)
        files = repairOrder.first().repair_images
        for i in files:
            file_path = i.replace(f'{settings.url_ROOT}/','')
            print(file_path)
            if os.path.exists(file_path):
                os.remove(file_path)
        report = repairOrder.first().reports
        if report.first():
            files = report.first().finished_images
            for i in files:
                file_path = i.replace(f'{settings.url_ROOT}/','')
                print(file_path)
                if os.path.exists(file_path):
                    os.remove(file_path)
        repairOrder.delete()
        # token = request.query_params.get('token')
        # page = request.query_params.get('page')
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/repairOrder_info/?{params}')


class repairInfoReport(APIView):
    def get(self,request,id):
        report = models.Report.objects.filter(id = id).first()
        repairID = report.repairOrder.id
        reportID = report.id
        reportInfo = report_info(instance=report)
        for field in reportInfo.fields.values():
            field.widget.attrs['readonly'] = True
            field.widget.attrs['disabled'] = True
            field.widget.attrs['style'] = 'background-color: #ffffff;'
        token = request.query_params.get('token')
        all_images = report.finished_images
        images = [url for url in all_images if not is_video(url)]
        print(images)
        videos = [url for url in all_images if is_video(url)]
        # page = request.query_params.get('page')
        # print(videos)
        params = request.META.get('QUERY_STRING', '')
        context = {'form': reportInfo,
                   'token': token,
                   'images': images,
                   'videos': videos,
                   "reportID":reportID,
                   'repairID': repairID,
                   # 'page':page
                   'params': params,
                   }
        return render(request, 'repairOrder/report_info_detail.html',context)
class repairOrderInfoQuotation(APIView):
    def get(self,request,id):
        token = request.query_params.get('token')
        repairOrder = models.RepairOrder.objects.get(id = id)
        quotation_status = repairOrder.quotation_status
        quotations = repairOrder.Quotations.all()
        data = QuotationSerializer(instance=quotations,many=True).data
        quotation = []
        filePath = ''
        if data:
            quotation = data[0]['quotation']
            filePath = data[0]['file_position']
        # page = request.query_params.get('page')
        params = request.META.get('QUERY_STRING', '')
        context = {
                   'all_quotations':quotation,
                   'quotation_status':quotation_status,
                   'filePath':filePath,
                   # 'page':page,
                    'params':params,
                    'token':token
                   }
        print(quotation)
        return render(request,'repairOrder/quotation.html',context)
    def post(self,request,id):
        data = request.data
        print(data)
        repairOrder = models.RepairOrder.objects.get(id = id)
        companyName = id
        fileName = repairOrder.order_id
        token = request.query_params.get('token')
        manager = models.ManagerInfo.objects.filter(token = token).first().name
        quotations = repairOrder.Quotations.first()
        fileData = []

        for item in data['quotations']:
            dataItem = []
            for key in item:
                dataItem.append(item[key])
            fileData.append(dataItem)
        print(fileData)
        price = 0
        if quotations:
            if data['quotations'] != []:
                quotations.quotation = data['quotations']
                price = data['quotations'][0]['tax_price']
                filePath = generate_quotation(companyName, fileName, fileData)
                quotations.file_position = filePath
                print(quotations.file_position)
                repairOrder.quotation_status = 1
            else:
                quotations.quotation = '保内无需报价'
                filePath = generate_quotation(companyName, fileName, fileData,status=1)
                repairOrder.quotation_status = 2
                quotations.file_position = filePath

            repairOrder.save()
            quotations.save()
        else:
            if data['quotations'] != []:
                quotationData = data['quotations']
                repairOrder.quotation_status = 1
                price = data['quotations'][0]['tax_price']
                filePath = generate_quotation(companyName, fileName, fileData)
            else:
                quotationData = '保内无需报价'
                repairOrder.quotation_status = 2
                filePath = generate_quotation(companyName, fileName, fileData, status=1)
            repairOrder.save()
            models.Quotation.objects.create(quotation = quotationData,repairOrder_id=id,responsible_manager=manager,file_position=filePath)

        openid = repairOrder.company.openid
        order_id = repairOrder.order_id
        machine = f'{repairOrder.machine_model}({repairOrder.machine_id})'
        template_id = "ry-dYB97Ra0VSEXfiHexCdcITtpexyCe1RnvraoDRU8"  # 派遣单模板
        page = f"/packageA/pages/repair_history/repair_history"
        data = {
            "character_string1": {"value": order_id},
            "thing2": {"value": machine},  # 根据实际模板调整
            "amount3": {"value": price},
            "time4": {"value": timezone.now().strftime("%Y-%m-%d")}
        }

        WeiChatAPI.send_subscribe_message(openid, template_id, page, data)

        return Response({'success':True})

class QuotationImport(APIView):
    def post(self,request):
        token = request.query_params.get('token')
        repairId = request.query_params.get('id')
        repairOrder = models.RepairOrder.objects.get(id = repairId)
        quotation = repairOrder.Quotations.first()
        manager = models.ManagerInfo.objects.filter(token=token).first().name
        if not quotation:
            quotation = models.Quotation.objects.create(responsible_manager=manager,repairOrder_id=repairId)

        if 'excel_file' not in request.FILES:
            return Response({'success': False, 'message': '请上传文件'}, status=400)
        try:
            excel_file = request.FILES['excel_file']

            # 读取Excel文件
            df = pd.read_excel(excel_file)

            # 验证必要列是否存在
            required_columns = ['名称及称号', '数量','含税金额']
            for col in required_columns:
                if col not in df.columns:
                    return Response({'success': False, 'message': f'缺少必要列: {col}'}, status=400)
            imported_count = 0
            quotations = []
            quotation_item = {}
            item_name = ['name','unit','quantity','rate','tax_unitprice','tax_price','tag']
            errors = []

            with transaction.atomic():
                for index, row in df.iterrows():
                    # 创建企业信息
                    # object = companyInfo.objects.filter(name=name[0])
                    # if not object:
                    try:
                        quotation_item['id'] = index+1
                        for i in range(len(df.columns)):
                            quotation_item[item_name[i]] = str(row.get(df.columns[i],''))
                        quotations.append(quotation_item.copy())
                        # print(quotations)
                        imported_count += 1
                    except Exception as e:
                        errors.append(f"第{index + 2}行错误: {str(e)}")
                        continue

            if errors:
                return Response({
                    'success': True,
                    'imported_count': imported_count,
                    'message': f'成功导入{imported_count}条，但有部分错误:\n' + '\n'.join(errors[:5])  # 只返回前5个错误
                })
            else:
                quotation.quotation = quotations
                quotation.save()
                return Response({
                    'success': True,
                    'imported_count': imported_count,
                    'message': f'成功导入{imported_count}条企业信息'
                })

        except Exception as e:
            print(e)
            return Response({
                'success': False,
                'message': f'处理文件时出错: {str(e)}'
            }, status=500)
class repairOrderInfoAddworker(APIView):
    def get(self,request,id):
        workersObjects = models.workerInfo.objects.all()
        workers = workerInfoSerializer(instance=workersObjects,many=True)
        token = request.query_params.get('token')
        repairOrder = models.RepairOrder.objects.filter(id = id).first()
        worker = repairOrder.worker
        # page = request.query_params.get('page')
        params = request.META.get('QUERY_STRING', '')
        if worker:
            workerID = worker.id
        else:
            workerID = 0
        return render(request,'repairOrder/AddWorker.html',{'params':params,'workers':workers.data,'workerID':workerID,'token':token})
    def post(self,request,id):
        repair_order = models.RepairOrder.objects.filter(id = id).first()
        data = request.data
        # print(data)
        repairDate = data.get('start_date')
        repair_order.worker_id = data['worker_id']
        repair_order.repair_date = repairDate
        repair_order.dispatch_status = 1
        repair_order.status = 1
        repair_order.save()

        #发送派遣消息
        openid = models.workerInfo.objects.filter(id = data['worker_id']).first().openid
        order_id = repair_order.order_id
        # contact_person = repair_order.responsible_name
        contact_person = "文臻荣"
        # contact_phone = repair_order.customer_phone
        contact_phone = 13823172206
        company = models.companyInfo.objects.filter(name = repair_order.customer_name).first()
        # adress = company.adress
        adress = "东关岸上林居"
        # repair_time = timezone.now().strftime('%Y-%m-%d')
        repair_time = repair_order.repair_date

        template_id="Q76qJL57pXYQnjpylIThYtRttxEm8FnfHAdCJhiwQeA" #派遣单模板
        page= f"pages/repair/detail?repairId={order_id}"
        data={
            "character_string6": {"value": order_id},
            "thing4": {"value": contact_person},  # 根据实际模板调整
            "phone_number5": {"value": contact_phone},
            "thing8":{"value":adress},
            "time12": {"value": repair_time}
            }

        WeiChatAPI.send_subscribe_message(openid,template_id,page,data)

        return Response({'success':True})



class reportConfirm(APIView):
    def get(self,request,id):
        report = models.Report.objects.filter(id = id).first()
        report.confirm_status = 1
        report.save()
        token = request.query_params.get('token')
        return redirect(f'/repair_info/{id}/report/?token={token}')


class adviceInfoView(APIView):
    def get(self, request):
        token = request.query_params.get('token')
        page = request.query_params.get('page')
        manager = authentication(token)
        params = request.META.get('QUERY_STRING', '')
        if not manager:
            # form.add_error("password","认证失败，请重新登录")
            error = '认证失败，请重新登录'
            return redirect(f'/?error={error}')
        search_dict = {}
        queryset = models.RepairAdvice.objects.all()
        # 1. 处理精确匹配的筛选条件
        exact_match_fields = {
            'satisfaction': 'satisfaction',
            'status': 'status',
        }

        for param, field in exact_match_fields.items():
            if value := request.GET.get(param):
                search_dict[field] = value

        # 2. 处理日期范围筛选
        if start_date := request.GET.get('start_date'):
            search_dict['created_at__gte'] = start_date
        if end_date := request.GET.get('end_date'):
            search_dict['created_at__lte'] = end_date

        # 3. 处理关键词模糊搜索
        if search := request.GET.get('search'):

            # 创建Q对象组合多个字段的模糊搜索
            q_objects = Q()
            search_fields = [
                'created_by__icontains',
                'customer_name__icontains'
            ]

            for field in search_fields:
                q_objects |= Q(**{field: search})

            # 将Q对象应用到queryset
            queryset = queryset.filter(q_objects)

            # 搜索外键worker
            # worker = models.workerInfo.objects.filter(name=search)
            # if worker:
            #     worker = models.workerInfo.objects.get(name=search)
            #     queryset = queryset.filter(**search_dict)

        # 4. 应用其他筛选条件
        if search_dict:
            queryset = queryset.filter(**search_dict)

        # 5. 排序（可选）
        queryset = queryset.order_by('-created_at')
        # queryset = models.RepairAdvice.objects.filter(**search_dict).order_by('id')
        # print(queryset)
        page_object = Pagination(request, queryset)
        queryset = page_object.page_queryset
        query = RepairAdviceSerializer(instance=queryset, many=True)
        if not page:
            page = 1
        # print(query.data)
        context = {
            # 'queryset': page_object.page_queryset,
            'queryset': query.data,
            'page_string': page_object.html(),
            'token':token,
            # 'page':page
            'params': params,
        }
        return render(request, 'repairAdvice/advice_info_list.html', context)
    def post(self,request):
        # token = request.query_params.get('token')
        # page = request.data.get('page')
        params = request.META.get('QUERY_STRING', '')
        # print(page)
        return redirect(f'/advice_info/?{params}')

class adviceInfoDetail(APIView):
    def get(self,request,id):
        advice = models.RepairAdvice.objects.filter(id = id).first()
        adviceInfo = advice_info(instance=advice)
        token = request.query_params.get('token')
        # page = request.query_params.get('page')
        all_images = advice.images
        params = request.META.get('QUERY_STRING', '')
        images = [url for url in all_images if not is_video(url)]
            # print(images)
        videos = [url for url in all_images if is_video(url)]
        for field in adviceInfo.fields.values():
            field.widget.attrs['readonly'] = True
            field.widget.attrs['disabled'] = True
            field.widget.attrs['style'] = 'background-color: #ffffff;'
        adviceID = advice.id
        adviceStatus = advice.status
        return render(request,'repairAdvice/advice_info_detail.html',
                      {'form':adviceInfo,'token':token,'images':images,'videos':videos,'params':params,
                       'adviceID':adviceID,'adviceStatus':adviceStatus})

class adviceInfoReply(APIView):
    def get(self,request,id):
        advice = models.RepairAdvice.objects.filter(id = id).first()
        adviceID = advice.id
        replyInfo = reply_info(instance=advice)
        # page = request.query_params.get('page')
        token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return render(request,'repairAdvice/reply_info.html',
                      {'form':replyInfo,'token':token,'params':params,'adviceID':adviceID})
    def post(self,request,id):
        advice = models.RepairAdvice.objects.filter(id=id).first()
        adviceID = advice.id
        data =request.data
        advice.reply = data.get('reply')
        date = timezone.now().strftime('%Y-%m-%d')
        advice.reply_date = date
        advice.status = 2
        advice.save()
        # page = request.query_params.get('page')
        # token = request.query_params.get('token')
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/advice_info/{adviceID}/reply/?{params}')

class adviceInfoDelete(APIView):
    def get(self,request,id):
        advice = models.RepairAdvice.objects.filter(id = id)
        files = advice.first().images
        if files:
            for file in files:
                file_path = file.replace(f'{settings.url_ROOT}/','')
                if os.path.exists(file_path):
                    os.remove(file_path)

        advice.delete()

        # token = request.query_params.get('token')
        # page = request.query_params.get('page')
        params = request.META.get('QUERY_STRING', '')
        return redirect(f'/advice_info/?token={params}')

from openpyxl.styles import Font
import openpyxl
# class workerTime(APIView):
#     authentication_classes = []
#
#     def get(self, request):
#         return render(request, 'attendanceInfo/attendanceInfo.html')
#
#     def post(self, request):
#         """
#         接收上传的考勤Excel/CSV文件，处理并计算每日工时，
#         然后将带有红色标记的结果以Excel文件的形式直接返回给用户下载。
#         """
#         if 'excel_file' not in request.FILES:
#             return Response({'success': False, 'message': '请上传文件，文件字段名为 "excel_file"'}, status=400)
#
#         excel_file = request.FILES['excel_file']
#
#         # 验证文件类型
#         name, ext = os.path.splitext(excel_file.name)
#         if ext.lower() not in ['.xlsx', '.csv']:
#             return Response({'success': False, 'message': '不支持的文件类型，请上传 .xlsx 或 .csv 文件。'}, status=400)
#
#         # 使用 FileSystemStorage 存储文件到 MEDIA_ROOT
#         fs = FileSystemStorage(location=settings.STATIC_ROOT)
#         filename = fs.save(excel_file.name, excel_file)
#         file_path = os.path.join(settings.STATIC_ROOT, filename)
#
#         try:
#             # 调用考勤处理函数，获取DataFrame和红色单元格信息
#             final_report_df, red_cell_info = clean_and_calculate_attendance(file_path, duplicate_threshold_minutes=5)
#
#             if final_report_df is None or final_report_df.empty:
#                 return Response({'success': False, 'message': '文件内容处理失败，请检查文件格式或数据完整性。'},
#                                 status=400)
#
#             # 将DataFrame保存到内存中的Excel文件
#             output = io.BytesIO()
#             with pd.ExcelWriter(output, engine='openpyxl') as writer:
#                 final_report_df.to_excel(writer, sheet_name='考勤报告', index=False)
#
#             # 重新打开工作簿设置红色字体
#             output.seek(0)
#             wb = openpyxl.load_workbook(output)
#             ws = wb.active
#
#             # 设置红色字体
#             red_font = Font(color="FF0000", bold=True)
#
#             # 找到"未打卡时间点"列的索引
#             column_index = None
#             for col_idx, cell in enumerate(ws[1], 1):
#                 if cell.value == "未打卡时间点":
#                     column_index = col_idx
#                     break
#
#             # 为需要设置红色的单元格设置样式
#             if column_index:
#                 for row_idx, text in red_cell_info:
#                     # Excel行索引从1开始，数据从第2行开始
#                     cell = ws.cell(row=row_idx + 2, column=column_index)
#                     if cell.value and any(missing_text in str(cell.value) for missing_text in [text, '未上岗']):
#                         cell.font = red_font
#
#             # 保存到BytesIO
#             final_output = io.BytesIO()
#             wb.save(final_output)
#             final_output.seek(0)
#
#             # 定义返回的文件名
#             response_filename = f"考勤工时报告_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
#
#             # 创建 FileResponse，直接返回文件给用户下载
#             response = FileResponse(
#                 final_output,
#                 content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
#             )
#             response['Content-Disposition'] = f'attachment; filename="{response_filename}"'
#
#             # 清理临时文件
#             if os.path.exists(file_path):
#                 os.remove(file_path)
#
#             return response
#
#         except Exception as e:
#             # 清理临时文件
#             if os.path.exists(file_path):
#                 os.remove(file_path)
#             return Response({'success': False, 'message': f'处理过程中发生错误: {str(e)}'}, status=500)
class workerTime(APIView):
    authentication_classes = []

    def get(self, request):
        return render(request, 'attendanceInfo/attendanceInfo.html')

    def post(self, request):
        """
        接收上传的考勤Excel/CSV文件，处理并返回结果。
        """
        if 'excel_file' not in request.FILES:
            return Response({'success': False, 'message': '请上传文件，文件字段名为 "excel_file"'}, status=400)

        excel_file = request.FILES['excel_file']

        # 验证文件类型并获取扩展名
        name, ext = os.path.splitext(excel_file.name)
        if ext.lower() not in ['.xlsx', '.csv', '.xls']:
            return Response({'success': False, 'message': '不支持的文件类型，请上传 .xlsx, .xls 或 .csv 文件。'},
                            status=400)

        try:
            # 直接将文件内容读入内存，创建文件流
            file_stream = io.BytesIO(excel_file.read())

            # 调用考勤处理函数，现在它需要文件流和文件扩展名
            final_report_df, red_cell_info = clean_and_calculate_attendance(file_stream, ext,
                                                                            duplicate_threshold_minutes=5)

            if final_report_df is None or final_report_df.empty:
                return Response({'success': False, 'message': '文件内容处理失败，请检查文件格式或数据完整性。'},
                                status=400)

            # 将DataFrame保存到内存中的Excel文件
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                final_report_df.to_excel(writer, sheet_name='考勤报告', index=False)

            # 重新打开工作簿设置红色字体
            output.seek(0)
            wb = openpyxl.load_workbook(output)
            ws = wb.active

            # 设置红色字体
            red_font = Font(color="FF0000", bold=True)

            column_index = None
            for col_idx, cell in enumerate(ws[1], 1):
                if cell.value == "未打卡时间点":
                    column_index = col_idx
                    break

            if column_index:
                for row_idx, text in red_cell_info:
                    cell = ws.cell(row=row_idx + 2, column=column_index)
                    if cell.value and any(missing_text in str(cell.value) for missing_text in [text, '未上岗']):
                        cell.font = red_font

            final_output = io.BytesIO()
            wb.save(final_output)
            final_output.seek(0)

            response_filename = f"考勤工时报告_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
            response = FileResponse(
                final_output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{response_filename}"'
            return response

        except Exception as e:
            return Response({'success': False, 'message': f'处理过程中发生错误: {str(e)}'}, status=500)






#营销系统部分
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, logout, authenticate
from django.core.exceptions import PermissionDenied
from django.http import JsonResponse
import json
from django.core.exceptions import ValidationError
from django.contrib import messages
from django.http import Http404
from django.db.models import Q
from datetime import datetime,timedelta,date
from .models import (
    CustomUser, Client, VisitRecord, ClientEquipment, Competitor,
    FollowUpPlan, AnnualPlan, MonthlyReport, WeeklyReport,ClientContact,
    ClientPurchase,ClientGeneration,SellsQuotation
)
from app01.serializers.sells import (
    ClientForm, VisitRecordForm, EquipmentForm,
    CompetitorForm, FollowUpForm, AnnualPlanForm,
    ContactForm, ClientGenerationForm,UserRegistrationForm,
    SalespersonEditForm,ClientPurchaseForm,SellsQuotationForm
)
from django.contrib.auth import get_user_model
from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import Group
# ==================用户注册模块 ====================
def register(request):
    # 如果用户已登录，重定向到首页
    if request.user.is_authenticated:
        return redirect('dashboard')
    if request.method == 'POST':
        # 强制设置用户类型为销售人员
        mutable_data = request.POST.copy()
        # mutable_data['user_type'] = 'sales'
        if mutable_data['user_type'] == 'admin':
            if mutable_data['admin_code'] != settings.SELLS_ADMIN_PASSWORD:
                messages.error(request, '管理员注册专用码错误！')
                # return redirect('register')
        form = UserRegistrationForm(mutable_data)

        if form.is_valid():
            user = form.save()
            messages.success(request, '注册成功！请登录')
            return redirect('login')
    else:
        form = UserRegistrationForm(initial={'user_type': 'sales'})

    return render(request, 'sells/auth/register.html', {'form': form})

# ==================== 用户认证模块 ====================
def user_login(request):
    if request.user.is_authenticated:
        return redirect('dashboard')

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')


        User = get_user_model()
        try:
            user = User.objects.get(username=username)
            if user.check_password(password):
                print("密码正确但authenticate失败，可能是认证后端问题")
            else:
                print("密码错误")
        except User.DoesNotExist:
            print("用户不存在")

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            next_url = request.GET.get('next', 'dashboard')
            return redirect(next_url)
        else:
            messages.error(request, '用户名或密码错误')

    return render(request, 'sells/auth/login.html')


@login_required
def user_logout(request):
    logout(request)
    return redirect('login')


# ==================== 仪表盘 ====================
@login_required
def dashboard(request):
    # 基础统计
    User = request.user
    if User.user_type == 'admin':
        stats = {
            'clients_count': Client.objects.count(),
            'visits_count': VisitRecord.objects.count(),
            'plans_count': AnnualPlan.objects.count(),
        }
    else:
        stats = {
            'clients_count': Client.objects.filter(assigned_salesperson=User).count(),
            'visits_count': VisitRecord.objects.filter(salesperson=User).count(),
            'plans_count': AnnualPlan.objects.count(),
        }

    # 最近动态
    if User.user_type == 'admin':
        recent_visits = VisitRecord.objects.order_by('-visit_date')[:5]
        upcoming_followups = FollowUpPlan.objects.filter(
            follow_up_date__gte=datetime.now().date()
        ).order_by('follow_up_date')[:5]
    else:
        recent_visits = VisitRecord.objects.filter(salesperson=User).order_by('-visit_date')[:5]
        upcoming_followups = FollowUpPlan.objects.filter(
            follow_up_date__gte=datetime.now().date(),responsible_person= User
        ).order_by('follow_up_date')[:5]

    # 销售数据概览
    if User.user_type == 'admin':
        monthly_reports = MonthlyReport.objects.order_by('-month')[:3]
    else:
        monthly_reports = MonthlyReport.objects.filter(
            reporter=User
        ).order_by('-month')[:3]
    if User.user_type == 'admin':
        weekly_reports = WeeklyReport.objects.order_by('-week_start')[:3]
    else:
        weekly_reports = WeeklyReport.objects.filter(
            reporter=User
        ).order_by('-week_start')[:3]

    return render(request, 'sells/dashboard.html', {
        'stats': stats,
        'recent_visits': recent_visits,
        'upcoming_followups': upcoming_followups,
        'monthly_reports': monthly_reports,
        'weekly_reports': weekly_reports,
    })


# ==================== 客户管理 ====================
@login_required
def client_list(request):
    # # 获取筛选参数
    #
    # query = request.GET.get('q', '')
    # clients = Client.objects.all()
    #
    # if query:
    #     clients = clients.filter(
    #         Q(name__icontains=query) |
    #         Q(contact_person__icontains=query) |
    #         Q(phone__icontains=query)
    #     )
    #
    # if request.user.user_type == 'sales':
    #     clients = clients.filter(assigned_salesperson=request.user)
    #
    # return render(request, 'sells/clients/list.html', {
    #     'clients': clients,
    #     'search_query': query,
    # })
    # 获取筛选参数
    name = request.GET.get('name', '').strip()
    contact_person = request.GET.get('contact_person', '').strip()
    phone = request.GET.get('phone', '').strip()
    industry = request.GET.get('industry', '').strip()

    # 构建查询条件
    clients = Client.objects.all()
    if request.user.user_type == 'sales':
        clients = clients.filter(assigned_salesperson=request.user)

    if name:
        clients = clients.filter(name__icontains=name)
    if contact_person:
        clients = clients.filter(contact_person__icontains=contact_person)
    if phone:
        clients = clients.filter(phone__icontains=phone)
    if industry:
        clients = clients.filter(industry=industry)

    # 排序和分页
    clients = clients.order_by('-created_at')
    paginator = Paginator(clients, 10)  # 每页10条
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'clients': page_obj,
        'page_obj': page_obj,
        'is_paginated': paginator.num_pages > 1,
        'page' : request.GET.get('page'),
    }

    return render(request, 'sells/clients/list.html', context)


@login_required
def client_detail(request, pk):
    client = get_object_or_404(Client, pk=pk)

    # 权限检查
    if request.user.user_type == 'sales' and client.assigned_salesperson != request.user:
        raise Http404("无权查看此客户")

    visits = VisitRecord.objects.filter(client=client).order_by('-visit_date')
    equipments = ClientEquipment.objects.filter(client=client)
    competitors = Competitor.objects.filter(client=client)
    contacts = ClientContact.objects.filter(client=client)
    generations = ClientGeneration.objects.filter(client=client)
    purchases = ClientPurchase.objects.filter(client=client)
    quotations = SellsQuotation.objects.filter(client=client)
    return render(request, 'sells/clients/detail.html', {
        'client': client,
        'visits': visits,
        'equipments': equipments,
        'competitors': competitors,
        'contacts': contacts,
        'generations' : generations,
        'purchases' : purchases,
        'quotations' : quotations,
        'page' : request.GET.get('page')
    })


@login_required
def client_create(request):
    if request.method == 'POST':
        form = ClientForm(request.POST)
        if form.is_valid():
            client = form.save(commit=False)
            client.created_by = request.user
            client.save()
            messages.success(request, f'客户 {client.name} 创建成功')
            return redirect('client_detail', pk=client.pk)
    else:
        form = ClientForm()

    return render(request, 'sells/form.html', {
        'form': form,
        'title': '创建新客户',
        'submit_text': '创建客户',
    })


@login_required
def client_update(request, pk):
    client = get_object_or_404(Client, pk=pk)

    # 权限检查
    if request.user.user_type != 'admin':
        raise Http404("无权修改客户信息")

    if request.method == 'POST':
        form = ClientForm(request.POST, instance=client)
        if form.is_valid():
            form.save()
            messages.success(request, f'客户 {client.name} 更新成功')
            return redirect('client_detail', pk=client.pk)
    else:
        form = ClientForm(instance=client)

    return render(request, 'sells/form.html', {
        'form': form,
        'title': f'编辑客户 - {client.name}',
        'submit_text': '更新信息',
    })


@login_required
def client_delete(request, pk):
    client = get_object_or_404(Client, pk=pk)

    # 权限检查
    if request.user.user_type != 'admin':
        raise Http404("无权删除客户")

    if request.method == 'POST':
        client_name = client.name
        client.delete()
        messages.success(request, f'客户 {client_name} 已删除')
        return redirect('client_list')

    return render(request, 'sells/confirm_delete.html', {
        'object': client,
        'title': '确认删除客户',
        'cancel_url': 'client_detail',
    })

#客户其他信息的删除

from django.views.decorators.http import require_POST
@require_POST
@login_required
def delete_contact(request, pk):
    try:
        contact = ClientContact.objects.get(pk=pk)
        contact.delete()
        return JsonResponse({'success': True})
    except ClientContact.DoesNotExist:
        return JsonResponse({'success': False, 'message': '联系人不存在'}, status=404)

# @require_POST
@login_required
def delete_equipment(request, pk):
    try:
        equipment = ClientEquipment.objects.get(pk=pk)
        equipment.delete()
        return JsonResponse({'success': True})
    except ClientEquipment.DoesNotExist:
        return JsonResponse({'success': False, 'message': '设备不存在'}, status=404)

@require_POST
@login_required
def delete_generation(request, pk):
    try:
        generation = ClientGeneration.objects.get(pk=pk)
        generation.delete()
        return JsonResponse({'success': True})
    except ClientGeneration.DoesNotExist:
        return JsonResponse({'success': False, 'message': '产品不存在'}, status=404)

@require_POST
@login_required
def delete_competitor(request, pk):
    try:
        competitor = Competitor.objects.get(pk=pk)
        competitor.delete()
        return JsonResponse({'success': True})
    except Competitor.DoesNotExist:
        return JsonResponse({'success': False, 'message': '竞争对手不存在'}, status=404)

@require_POST
@login_required
def delete_purchase(request, pk):
    try:
        purchase = ClientPurchase.objects.get(pk=pk)
        purchase.delete()
        return JsonResponse({'success': True})
    except ClientPurchase.DoesNotExist:
        return JsonResponse({'success': False, 'message': '购买记录不存在'}, status=404)

@require_POST
@login_required
def delete_quotation(request, pk):
    try:
        quotation = SellsQuotation.objects.get(pk=pk)
        quotation.delete()
        return JsonResponse({'success': True})
    except SellsQuotation.DoesNotExist:
        return JsonResponse({'success': False, 'message': '报价不存在'}, status=404)


# ==================== 拜访记录 ====================
@login_required
def visit_list(request):
    salesperson_id = request.GET.get('salesperson')
    client_name = request.GET.get('client')
    visit_date = request.GET.get('visit_date')

    # 获取所有销售人员和客户用于下拉菜单
    salespersons = CustomUser.objects.filter(user_type = "sales").order_by('username')
    clients = Client.objects.all().order_by('name')
    search_client = Client.objects.filter(name = client_name).first()

    # 构建查询条件
    visits = VisitRecord.objects.select_related('client', 'salesperson').all()
    if client_name:
        # visits = visits.filter(client_id=client_id)
        visits = search_client.visit_records
    if salesperson_id:
        visits = visits.filter(salesperson_id=salesperson_id)
    if visit_date:
        visits = visits.filter(visit_date=visit_date)

    # 排序和分页
    visits = visits.order_by('-visit_date')
    paginator = Paginator(visits, 10)  # 每页10条记录
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'visits': page_obj,
        'page_obj': page_obj,
        'is_paginated': paginator.num_pages > 1,
        'salespersons': salespersons,
        'clients': clients,
        'page' : request.GET.get('page'),
    }

    return render(request, 'sells/visits/list.html', context)


@login_required
def visit_detail(request, pk):
    visit = get_object_or_404(VisitRecord, pk=pk)

    # 权限检查
    if request.user.user_type == 'sales' and visit.salesperson != request.user:
        raise Http404("无权查看此拜访记录")

    followups = FollowUpPlan.objects.filter(visit_record=visit)

    return render(request, 'sells/visits/detail.html', {
        'visit': visit,
        'followups': followups,
        'page':request.GET.get('page')
    })


@login_required
def visit_create(request):
    if request.method == 'POST':
        form = VisitRecordForm(request.POST, user=request.user)
        if form.is_valid():
            visit = form.save(commit=False)
            visit.salesperson = request.user
            visit.save()
            messages.success(request, '拜访记录创建成功')
            return redirect('visit_detail', pk=visit.pk)
    else:
        initial = {}
        if 'client_id' in request.GET:
            initial['client'] = request.GET['client_id']
        form = VisitRecordForm(user=request.user, initial=initial)

    return render(request, 'sells/form.html', {
        'form': form,
        'title': '新建拜访记录',
        'submit_text': '保存记录',
    })



@login_required
def visit_record_detail(request, record_id):
    """
    拜访记录详情视图
    """
    record = get_object_or_404(VisitRecord, record_id=record_id)

    # 准备上下文数据
    context = {
        'record': record,
        'purpose_display': record.get_purpose_display(),
        'method_display': record.get_method_display(),
        'intention_display': record.get_cooperation_intention_display(),
        'page' : request.GET.get('page')
    }

    return render(request, 'sells/visits/visit_record_detail.html', context)


@login_required
def visit_update(request, pk):
    visit = get_object_or_404(VisitRecord, pk=pk)

    # 权限检查
    if request.user.user_type == 'sales' and visit.salesperson != request.user:
        raise Http404("无权修改此拜访记录")

    if request.method == 'POST':
        form = VisitRecordForm(request.POST, instance=visit, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, '拜访记录更新成功')
            return redirect('visit_detail', pk=visit.pk)
    else:
        form = VisitRecordForm(instance=visit, user=request.user)

    return render(request, 'sells/form.html', {
        'form': form,
        'title': '编辑拜访记录',
        'submit_text': '更新记录',
    })



from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import VisitRecord


@login_required
def visit_delete(request, pk):
    # 获取要删除的记录
    record = get_object_or_404(VisitRecord, pk=pk)

    # 检查当前用户是否有权限删除该记录
    if request.user.user_type != "admin" and not record.salesperson != request.user:
        Http404(request, '您没有权限删除此拜访记录')
        return redirect('visit_list')  # 假设有一个拜访记录列表的URL

    if request.method == 'POST':
        # 执行删除操作
        client_name = record.client.name
        record.delete()
        messages.success(request, f'成功删除客户 {client_name} 的拜访记录')
        return redirect('visit_list')

    # 如果不是POST请求，重定向到列表页
    return render(request, 'sells/visits/confirm_delete.html', {'record':record})


# ==================== 年度计划 ====================
@login_required
def annual_plan_list(request):
    plans = AnnualPlan.objects.all().order_by('-year')
    return render(request, 'sells/plans/annual_list.html', {
        'plans': plans,
    })


@login_required
def annual_plan_detail(request, pk):
    plan = get_object_or_404(AnnualPlan, pk=pk)
    return render(request, 'sells/plans/annual_detail.html', {
        'plan': plan,
    })


@login_required
def annual_plan_create(request):
    if request.user.user_type != 'admin':
        raise Http404("无权创建年度计划")

    if request.method == 'POST':
        form = AnnualPlanForm(request.POST)
        if form.is_valid():
            plan = form.save(commit=False)
            plan.created_by = request.user
            plan.save()
            messages.success(request, f'{plan.year}年度计划创建成功')
            return redirect('annual_plan_detail', pk=plan.pk)
    else:
        form = AnnualPlanForm(initial={'year': datetime.now().year})

    return render(request, 'sells/form.html', {
        'form': form,
        'title': '创建年度销售计划',
        'submit_text': '创建计划',
    })


@login_required
def annual_plan_update(request, pk):
    plan = get_object_or_404(AnnualPlan, pk=pk)

    # 权限检查：只有管理员或创建者可以修改
    if not (request.user.user_type == 'admin' or plan.created_by == request.user):
        raise PermissionDenied("无权修改此计划")

    if request.method == 'POST':
        form = AnnualPlanForm(request.POST, instance=plan)
        if form.is_valid():
            form.save()
            messages.success(request, f'{plan.year}年度计划更新成功')
            return redirect('annual_plan_detail', pk=plan.pk)
    else:
        form = AnnualPlanForm(instance=plan)

    return render(request, 'sells/plans/annual_form.html', {
        'form': form,
        'title': f'编辑 {plan.year} 年度计划',
        'plan': plan
    })


@login_required
def annual_plan_delete(request, pk):
    plan = get_object_or_404(AnnualPlan, pk=pk)

    # 权限检查：只有管理员或创建者可以删除
    if not (request.user.user_type == 'admin' or plan.created_by == request.user):
        raise PermissionDenied("无权删除此计划")

    if request.method == 'POST':
        year = plan.year
        plan.delete()
        messages.success(request, f'{year}年度计划已删除')
        return redirect('annual_plan_list')

    return render(request, 'sells/plans/annual_confirm_delete.html', {
        'plan': plan
    })


# ==================== 报表管理 ====================
@login_required
def monthly_report_create(request):
    if request.method == 'POST':
        try:
            # 解析JSON数据
            data = json.loads(request.body)

            # 验证必填字段
            required_fields = [
                'month', 'team', 'monthly_sales', 'equipment_sold',
                'sales_analysis', 'new_clients_count', 'existing_clients_visited',
                'existing_client_feedback', 'repeat_purchase_amount'
            ]

            # for field in required_fields:
            #     if not data.get(field):
            #         return JsonResponse({'status': 'error', 'message': f'字段 {field} 是必填项'}, status=400)

            # 处理JSON字段
            json_fields = [
                'equipment_details', 'new_clients_details',
                'new_client_challenges', 'repeat_purchase_details'
            ]

            for field in json_fields:
                if field in data and isinstance(data[field], str):
                    try:
                        data[field] = json.loads(data[field])
                    except json.JSONDecodeError:
                        data[field] = {}
            #处理月份格式
            month_str = data['month']
            if len(month_str) == 7:  # 格式为 YYYY-MM
                month_str += "-01"

                # 验证日期格式
            month_date = datetime.strptime(month_str, "%Y-%m-%d").date()
            # 创建月度报表
            report = MonthlyReport(
                reporter=request.user,
                month=month_date,
                team=data['team'],
                monthly_sales=data['monthly_sales'],
                equipment_sold=data['equipment_sold'],
                equipment_details=data.get('equipment_details', {}),
                sales_analysis=data['sales_analysis'],
                new_clients_count=data['new_clients_count'],
                new_clients_details=data.get('new_clients_details', {}),
                new_client_challenges=data.get('new_client_challenges', {}),
                existing_clients_visited=data['existing_clients_visited'],
                existing_client_feedback=data['existing_client_feedback'],
                repeat_purchase_amount=data['repeat_purchase_amount'],
                repeat_purchase_details=data.get('repeat_purchase_details', {}),
                promotion_activities=data.get('promotion_activities', ''),
                strategy_implementation=data.get('strategy_implementation', ''),
                industry_trends=data.get('industry_trends', ''),
                competitor_analysis=data.get('competitor_analysis', ''),
                customer_demand_changes=data.get('customer_demand_changes', ''),
                training_progress=data.get('training_progress', ''),
                skill_improvements=data.get('skill_improvements', ''),
                product_challenges=data.get('product_challenges', ''),
                market_challenges=data.get('market_challenges', ''),
                customer_challenges=data.get('customer_challenges', ''),
                personal_challenges=data.get('personal_challenges', ''),
                challenge_solutions=data.get('challenge_solutions', ''),
                next_month_sales_target=data.get('next_month_sales_target', 0),
                next_month_equipment_target=data.get('next_month_equipment_target', 0),
                next_month_new_clients_target=data.get('next_month_new_clients_target', 0),
                next_month_repeat_purchase_target=data.get('next_month_repeat_purchase_target', 0),
                key_tasks=data.get('key_tasks', ''),
                hr_needs=data.get('hr_needs', ''),
                financial_needs=data.get('financial_needs', ''),
                technical_needs=data.get('technical_needs', '')
            )

            # 验证并保存
            # report.full_clean()
            report.save()

            return JsonResponse({'status': 'success', 'message': '报表提交成功'})

        except ValidationError as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    else:
        return render(request, 'sells/reports/monthly/monthly_form.html')

    # return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)




@login_required
def monthly_report_detail(request, pk):
    report = get_object_or_404(MonthlyReport, pk=pk)

    # 权限检查：只有管理员或创建者可以查看
    if not (request.user.user_type == 'admin' or report.reporter == request.user):
        raise PermissionDenied("无权查看此报表")

    # 处理JSON字段
    # equipment_details = report.equipment_details or {}
    # new_clients_details = report.new_clients_details or []
    # new_client_challenges = report.new_client_challenges or []
    # repeat_purchase_details = report.repeat_purchase_details or []

    context = {
        'report': report,
        'page_title': f"{report.month:'Y年m月'}报表详情"
    }


    # return render(request, 'sells/reports/monthly_report_detail.html', {
    #     'report': report,
    #     'equipment_details': equipment_details,
    #     'new_clients_details': new_clients_details,
    #     'new_client_challenges': new_client_challenges,
    #     'repeat_purchase_details': repeat_purchase_details,
    # })
    return render(request, 'sells/reports/monthly/monthly_report_detail.html', context)




@login_required
def monthly_report_edit(request, pk):
    report = get_object_or_404(MonthlyReport, pk=pk)

    # 权限检查 - 只有创建者或管理员可以编辑
    if not request.user.is_superuser and report.reporter != request.user:
        return JsonResponse({'status': 'error', 'message': '无权编辑此报表'}, status=403)

    if request.method == 'GET':
        return render(request, 'sells/reports/monthly/monthly_report_edit.html', {'report': report})

    elif request.method == 'POST':
        try:
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            month_str = data['month']
            if len(month_str) == 7:  # 格式为 YYYY-MM
                month_str += "-01"

                # 验证日期格式
            month_date = datetime.strptime(month_str, "%Y-%m-%d").date()
            # 更新报表数据
            report.month = month_date
            report.team = data.get('team', report.team)
            report.monthly_sales = data.get('monthly_sales', report.monthly_sales)
            report.equipment_sold = data.get('equipment_sold', report.equipment_sold)
            report.equipment_details = data.get('equipment_details', report.equipment_details)
            report.sales_analysis = data.get('sales_analysis', report.sales_analysis)

            # 更新其他字段...
            report.new_clients_count = data.get('new_clients_count', report.new_clients_count)
            report.new_clients_details = data.get('new_clients_details', report.new_clients_details)
            report.existing_clients_visited = data.get('existing_clients_visited', report.existing_clients_visited)
            report.existing_client_feedback = data.get('existing_client_feedback', report.existing_client_feedback)
            report.repeat_purchase_amount = data.get('repeat_purchase_amount', report.repeat_purchase_amount)
            report.repeat_purchase_details = data.get('repeat_purchase_details', report.repeat_purchase_details)

            # 更新问题和解决方案...
            report.product_challenges = data.get('product_challenges', report.product_challenges)
            report.market_challenges = data.get('market_challenges', report.market_challenges)
            report.customer_challenges = data.get('customer_challenges', report.customer_challenges)
            report.personal_challenges = data.get('personal_challenges', report.personal_challenges)
            report.challenge_solutions = data.get('challenge_solutions', report.challenge_solutions)

            # 更新下月计划...
            report.next_month_sales_target = data.get('next_month_sales_target', report.next_month_sales_target)
            report.next_month_equipment_target = data.get('next_month_equipment_target',
                                                          report.next_month_equipment_target)
            report.next_month_new_clients_target = data.get('next_month_new_clients_target',
                                                            report.next_month_new_clients_target)
            report.next_month_repeat_purchase_target = data.get('next_month_repeat_purchase_target',
                                                                report.next_month_repeat_purchase_target)
            report.key_tasks = data.get('key_tasks', report.key_tasks)
            report.hr_needs = data.get('hr_needs', report.hr_needs)
            report.financial_needs = data.get('financial_needs', report.financial_needs)
            report.technical_needs = data.get('technical_needs', report.technical_needs)

            report.save()

            if request.content_type == 'application/json':
                return JsonResponse({'status': 'success'})
            return redirect(report)

        except Exception as e:
            if request.content_type == 'application/json':
                return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
            return render(request, 'sells/reports/monthly/monthly_report_edit.html', {
                'report': report,
                'error': str(e)
            })


@login_required
def monthly_report_delete(request, pk):
    report = get_object_or_404(MonthlyReport, pk=pk)

    # 权限检查：只有管理员或创建者可以删除
    if not (request.user.user_type == 'admin' or report.reporter == request.user):
        raise PermissionDenied("无权删除此报表")

    if request.method == 'POST':
        month = report.month.strftime('%Y年%m月')
        report.delete()
        messages.success(request, f'{month}月度报表已删除')
        return redirect('dashboard')

    return render(request, 'sells/reports/monthly/monthly_confirm_delete.html', {
        'report': report
    })


@login_required
def weekly_report_create(request):
    if request.method == 'POST':
        # form = WeeklyReportForm(request.POST)
        # if form.is_valid():
        #     report = form.save(commit=False)
        #     report.reporter = request.user
        #     report.save()
        #     messages.success(request, '周报提交成功')
        #     return redirect('dashboard')
        try:
            try:
                data = json.loads(request.body)
            except json.JSONDecodeError as e:
                return JsonResponse({
                    'status': 'error',
                    'message': f'无效的JSON数据: {str(e)}'
                }, status=400)

            # 创建周报表
            report = WeeklyReport(
                week_start=data['week_start'],
                week_end=data['week_end'],
                reporter=request.user,
                team=data['team'],
                weekly_sales=data['weekly_sales'],
                last_week_sales=data['last_week_sales'],
                weekly_target=data['weekly_target'],
                equipment_sold=data['equipment_sold'],
                equipment_details=data.get('equipment_details', {}),
                new_clients_count=data['new_clients_count'],
                new_clients_details=data.get('new_clients_details', {}),
                existing_clients_visited=data['existing_clients_visited'],
                next_week_sales_target=data['next_week_sales_target'],
                next_week_equipment_target=data['next_week_equipment_target'],
                next_week_new_clients_target=data['next_week_new_clients_target'],
                next_week_repeat_purchase_target=data['next_week_repeat_purchase_target']
            )
            report.save()

            return JsonResponse({
                'status': 'success',
                'redirect_url': f'/reports/weekly/{report.report_id}/'
            })

        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=400)
    else:
        # today = datetime.now().date()
        # start = today - timedelta(days=today.weekday())
        # end = start + timedelta(days=6)
        #
        # form = WeeklyReportForm(initial={
        #     'week_start': start,
        #     'week_end': end,
        #     'team': request.user.groups.first().name if request.user.groups.exists() else ''
        # })
        return render(request, 'sells/reports/weekly/weekly_form.html')

    # return render(request, 'sells/reports/weekly_form.html', {
    #     'form': form,
    #     'title': '创建周报',
    # })




def weekly_report_detail(request, pk):
    report = get_object_or_404(WeeklyReport, pk=pk)

    # 权限检查（可选）
    if not request.user.is_superuser and report.reporter != request.user:
        return Http404("您没有权限查看此报表")

    context = {
        'report': report,
        'page_title': f"周报表详情 - {report.week_start:'Y-m-d'}"
    }

    return render(request, 'sells/reports/weekly/weekly_report_detail.html', context)


@login_required
def weekly_report_edit(request, pk):
    report = get_object_or_404(WeeklyReport, pk=pk)

    # 权限检查 - 只有创建者或管理员可以编辑
    if not request.user.is_superuser and report.reporter != request.user:
        return JsonResponse({'status': 'error', 'message': '无权编辑此报表'}, status=403)

    if request.method == 'GET':
        return render(request, 'sells/reports/weekly/weekly_report_edit.html', {'report': report})

    elif request.method == 'POST':
        try:
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST

            # 更新报表数据
            report.week_start = data.get('week_start', report.week_start)
            report.week_end = data.get('week_end', report.week_end)
            report.team = data.get('team', report.team)
            report.weekly_sales = data.get('weekly_sales', report.weekly_sales)
            report.last_week_sales = data.get('last_week_sales', report.last_week_sales)
            report.weekly_target = data.get('weekly_target', report.weekly_target)
            report.equipment_sold = data.get('equipment_sold', report.equipment_sold)
            report.equipment_details = data.get('equipment_details', report.equipment_details)
            report.new_clients_count = data.get('new_clients_count', report.new_clients_count)
            report.new_clients_details = data.get('new_clients_details', report.new_clients_details)
            report.existing_clients_visited = data.get('existing_clients_visited', report.existing_clients_visited)
            report.next_week_sales_target = data.get('next_week_sales_target', report.next_week_sales_target)
            report.next_week_equipment_target = data.get('next_week_equipment_target',
                                                         report.next_week_equipment_target)
            report.next_week_new_clients_target = data.get('next_week_new_clients_target',
                                                           report.next_week_new_clients_target)
            report.next_week_repeat_purchase_target = data.get('next_week_repeat_purchase_target',
                                                               report.next_week_repeat_purchase_target)

            report.save()

            if request.content_type == 'application/json':
                return JsonResponse({'status': 'success',})
            return redirect('weekly_report_detail', pk=report.report_id)

        except Exception as e:
            if request.content_type == 'application/json':
                return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
            return render(request, 'sells/reports/weekly/weekly_report_edit.html', {
                'report': report,
                'error': str(e)
            })




@login_required
def weekly_report_delete(request, pk):
    report = get_object_or_404(WeeklyReport, pk=pk)

    # 权限检查：只有报表创建者或管理员可以删除
    if not (request.user.is_superuser or report.reporter == request.user):
        messages.error(request, "您没有权限删除此报表")
        return redirect('weekly_report_detail', pk=pk)

    if request.method == 'POST':
        try:
            report.delete()
            messages.success(request, f"已成功删除 {report.week_start:'Y-m-d'} 的周报表")
            return redirect('dashboard')
        except Exception as e:
            messages.error(request, f"删除失败: {str(e)}")
            return redirect('weekly_report_detail', pk=pk)

    # GET请求显示确认页面
    context = {
        'report': report,
        'page_title': '确认删除周报表'
    }
    return render(request, 'sells/reports/weekly/weekly_report_confirm_delete.html', context)

# ==================== 销售人员管理 ====================



def is_admin(user):
    return user.user_type == 'admin'


@login_required
@user_passes_test(is_admin, login_url='/')
def salesperson_list(request):
    # 获取筛选参数
    username = request.GET.get('username', '')
    group_id = request.GET.get('group', '')

    # 获取所有销售人员
    salespersons = CustomUser.objects.filter(user_type='sales')

    # 应用筛选
    if username:
        salespersons = salespersons.filter(Q(username__icontains=username) | Q(user_name__icontains=username))

    if group_id:
        salespersons = salespersons.filter(groups__id=group_id)

    # 获取所有用户组用于筛选
    groups = Group.objects.all()

    context = {
        'salespersons': salespersons,
        'groups': groups,
        'search_username': username,
        'selected_group': group_id,
    }
    return render(request, 'sells/salesperson/list.html', context)


@login_required
@user_passes_test(is_admin, login_url='/')
def salesperson_detail(request, pk):
    salesperson = get_object_or_404(CustomUser, id=pk, user_type='sales')
    return render(request, 'sells/salesperson/detail.html', {'salesperson': salesperson})


@login_required
@user_passes_test(is_admin, login_url='/')
def edit_salesperson(request, pk):
    salesperson = get_object_or_404(CustomUser, id=pk, user_type='sales')

    if request.method == 'POST':
        form = SalespersonEditForm(request.POST, instance=salesperson)
        if form.is_valid():
            user = form.save()
            messages.success(request, f'成功更新销售人员: {user.username}')
            return redirect('salesperson_detail', pk=user.id)
    else:
        form = SalespersonEditForm(instance=salesperson)

    return render(request, 'sells/salesperson/edit_salesperson.html', {
        'form': form,
        'salesperson': salesperson
    })


from django.contrib.auth.hashers import make_password

#重置密码
# @login_required
# @user_passes_test(is_admin, login_url='/')
# def reset_salesperson_password(request, pk):
#     salesperson = get_object_or_404(CustomUser, id=pk, user_type='sales')
#
#     if request.method == 'POST':
#         # 生成随机密码
#         new_password = '88888888'
#         salesperson.password = make_password(new_password)
#         salesperson.save()
#
#         messages.success(request, f'已重置 {salesperson.username} 的密码')
#         return redirect('edit_salesperson', pk=salesperson.id)
#
#     # 确认页面
#     return render(request, 'sells/salesperson/confirm_password_reset.html', {'salesperson': salesperson})

@login_required
@user_passes_test(is_admin, login_url='/')
def reset_salesperson_password(request, pk):
    salesperson = get_object_or_404(CustomUser, id=pk, user_type='sales')

    if request.method == 'POST':
        # 生成随机密码
        new_password = '88888888'
        salesperson.password = make_password(new_password)
        salesperson.save()

        # 这里应该添加发送邮件给用户的逻辑
        # send_password_reset_email(salesperson.email, new_password)

        # messages.success(request, f'已重置 {salesperson.username} 的密码')

        # 返回JSON响应用于弹窗
        if request.content_type == 'application/json':
            return JsonResponse({
                'status': 'success',
                'message': f'成功重置 {salesperson.username} 的密码',
            })

        return redirect('edit_salesperson', pk=salesperson.id)

    # 如果不是POST请求，返回确认页面
    return render(request, 'sells/salesperson/confirm_password_reset.html', {'salesperson': salesperson})

@login_required
@user_passes_test(is_admin, login_url='/')
def delete_salesperson(request, pk):
    if request.method == 'POST':
        salesperson = get_object_or_404(CustomUser, id=pk, user_type='sales')
        username = salesperson.username
        salesperson.delete()
        messages.success(request, f'成功删除销售人员: {username}')
        return redirect('salesperson_list')
    else:
        return redirect('salesperson_list')


# ==================== 用户组管理 ====================
from django.db.models import Count
@login_required
@user_passes_test(is_admin)
def group_list(request):
    groups = Group.objects.all().order_by('name')
    return render(request, 'sells/groups/group_list.html', {'groups': groups})


@login_required
@user_passes_test(is_admin)
def group_add(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        if name:
            group, created = Group.objects.get_or_create(name=name)
            if created:
                messages.success(request, f'用户组 "{name}" 添加成功')
                return redirect('group_list')
            else:
                messages.error(request, f'用户组 "{name}" 已存在')
        else:
            messages.error(request, '组名不能为空')

    return render(request, 'sells/groups/group_form.html', {
        'title': '添加用户组',
        'submit_text': '添加',
        'cancel_url': 'group_list'
    })


@login_required
@user_passes_test(is_admin)
def group_edit(request, group_id):
    group = get_object_or_404(Group, pk=group_id)

    if request.method == 'POST':
        name = request.POST.get('name')
        if name and name != group.name:
            if not Group.objects.filter(name=name).exists():
                group.name = name
                group.save()
                messages.success(request, '用户组更新成功')
                return redirect('group_list')
            else:
                messages.error(request, f'用户组 "{name}" 已存在')
        else:
            messages.error(request, '组名不能为空或未更改')

    return render(request, 'sells/groups/group_form.html', {
        'title': f'编辑用户组: {group.name}',
        'submit_text': '更新',
        'cancel_url': 'group_list',
        'group': group
    })


@login_required
@user_passes_test(is_admin)
def group_delete(request, group_id):
    group = get_object_or_404(Group, pk=group_id)

    if request.method == 'POST':
        group_name = group.name
        group.delete()
        messages.success(request, f'用户组 "{group_name}" 已删除')
        return redirect('group_list')

    return render(request, 'sells/groups/group_confirm_delete.html', {'group': group})





from django.views.generic import ListView





# ==================== 其他功能 ====================
@login_required
def add_equipment(request, client_id):
    client = get_object_or_404(Client, pk=client_id)

    if request.method == 'POST':
        form = EquipmentForm(request.POST)
        if form.is_valid():
            equipment = form.save(commit=False)
            equipment.client = client
            equipment.save()
            messages.success(request, '设备信息添加成功')
            return redirect('client_detail', pk=client.pk)
    else:
        form = EquipmentForm()

    return render(request, 'sells/form.html', {
        'form': form,
        'title': f'为 {client.name} 添加设备',
        'submit_text': '添加设备',
    })

def update_equipment(request, pk):
    equipment = get_object_or_404(ClientEquipment, pk=pk)
    if request.method == 'POST':
        form = EquipmentForm(request.POST, instance=equipment)
        if form.is_valid():
            form.save()
            messages.success(request, '设备更新成功')
            return redirect('client_detail', pk=equipment.client.pk)
    else:
        form = EquipmentForm(instance=equipment)
    context = {
        'form': form,
        'title': '更新设备',
        'submit_text': '更新设备'
    }
    return render(request, 'sells/form.html', context)

@login_required
def add_contact(request, client_id):
    client = get_object_or_404(Client, pk=client_id)

    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            contact = form.save(commit=False)
            contact.client = client
            contact.save()
            messages.success(request, '联系人信息添加成功')
            return redirect('client_detail', pk=client.pk)
    else:
        form = ContactForm()

    return render(request, 'sells/form.html', {
        'form': form,
        'title': f'为 {client.name} 添加联系人',
        'submit_text': '添加联系人',
    })

def update_contact(request, pk):
    contact = get_object_or_404(ClientContact, pk=pk)
    if request.method == 'POST':
        form = ContactForm(request.POST, instance=contact)
        if form.is_valid():
            form.save()
            messages.success(request, '联系方式更新成功')
            return redirect('client_detail', pk=contact.client.pk)
    else:
        form = ContactForm(instance=contact)
    context = {
        'form': form,
        'title': '更新联系方式',
        'submit_text': '更新联系方式'
    }
    return render(request, 'sells/form.html', context)


@login_required
def add_competitor(request, client_id):
    client = get_object_or_404(Client, pk=client_id)

    if request.method == 'POST':
        form = CompetitorForm(request.POST)
        if form.is_valid():
            competitor = form.save(commit=False)
            competitor.client = client
            competitor.save()
            messages.success(request, '竞争对手信息添加成功')
            return redirect('client_detail', pk=client.pk)
    else:
        form = CompetitorForm()

    return render(request, 'sells/form.html', {
        'form': form,
        'title': f'为 {client.name} 添加竞争对手',
        'submit_text': '添加信息',
    })

def update_competitor(request, pk):
    competitor = get_object_or_404(Competitor, pk=pk)
    if request.method == 'POST':
        form = CompetitorForm(request.POST, instance=competitor)
        if form.is_valid():
            form.save()
            messages.success(request, '设备更新成功')
            return redirect('client_detail', pk=competitor.client.pk)
    else:
        form = CompetitorForm(instance=competitor)
    context = {
        'form': form,
        'title': '更新竞争对手',
        'submit_text': '更新竞争对手'
    }
    return render(request, 'sells/form.html', context)

@login_required
def add_followup(request, visit_id):
    visit = get_object_or_404(VisitRecord, pk=visit_id)

    if request.method == 'POST':
        form = FollowUpForm(request.POST)
        if form.is_valid():
            followup = form.save(commit=False)
            followup.visit_record = visit
            followup.save()
            messages.success(request, '跟进计划添加成功')
            return redirect('visit_detail', pk=visit.pk)
    else:
        form = FollowUpForm(initial={
            'responsible_person': request.user,
            'follow_up_date': datetime.now().date() + timedelta(days=7)
        })

    return render(request, 'sells/form.html', {
        'form': form,
        'title': f'为 {visit.client.name} 拜访添加跟进',
        'submit_text': '添加跟进',
    })



def add_client_generation(request, client_id):
    client = get_object_or_404(Client, pk=client_id)

    if request.method == 'POST':
        form = ClientGenerationForm(request.POST)
        if form.is_valid():
            generation = form.save(commit=False)
            generation.client = client
            generation.save()
            messages.success(request, '产品添加成功！')
            return redirect('client_detail', pk=client_id)  # 假设有一个客户详情页
    else:
        form = ClientGenerationForm()

    context = {
        'form': form,
        'title': f'为 {client.name} 添加所做产品',
        'submit_text': '添加客户产品'
    }
    return render(request, 'sells/form.html', context)

def update_generation(request, pk):
    generation = get_object_or_404(ClientGeneration, pk=pk)
    if request.method == 'POST':
        form = ClientGenerationForm(request.POST, instance=generation)
        if form.is_valid():
            form.save()
            messages.success(request, '设备更新成功')
            return redirect('client_detail', pk=generation.client.pk)
    else:
        form = ClientGenerationForm(instance=generation)
    context = {
        'form': form,
        'title': '更新所做产品',
        'submit_text': '更新所做产品'
    }
    return render(request, 'sells/form.html', context)

def add_client_purchase(request, client_id):
    client = get_object_or_404(Client, pk=client_id)

    if request.method == 'POST':
        form = ClientPurchaseForm(request.POST)
        if form.is_valid():
            purchase = form.save(commit=False)
            purchase.client = client
            purchase.save()
            messages.success(request, '购买记录添加成功！')
            return redirect('client_detail', pk=client_id)  # 假设有一个客户详情页
    else:
        form = ClientPurchaseForm()

    context = {
        'form': form,
        'title': f'为 {client.name} 添加购买记录',
        'submit_text': '添加购买记录'
    }
    return render(request, 'sells/form.html', context)

def update_client_purchase(request, pk):
    purchase = get_object_or_404(ClientPurchase, pk=pk)
    if request.method == 'POST':
        form = ClientPurchaseForm(request.POST, instance=purchase)
        if form.is_valid():
            form.save()
            messages.success(request, '设备更新成功')
            return redirect('client_detail', pk=purchase.client.pk)
    else:
        form = ClientPurchaseForm(instance=purchase)
    context = {
        'form': form,
        'title': '更新购买记录',
        'submit_text': '更新购买记录'
    }
    return render(request, 'sells/form.html', context)

def add_client_quotation(request, client_id):
    client = get_object_or_404(Client, pk=client_id)

    if request.method == 'POST':
        form = SellsQuotationForm(request.POST)
        if form.is_valid():
            quotation = form.save(commit=False)
            quotation.client = client
            quotation.save()
            messages.success(request, '报价添加成功！')
            return redirect('client_detail', pk=client_id)  # 假设有一个客户详情页
    else:
        form = SellsQuotationForm()

    context = {
        'form': form,
        'title': f'为 {client.name} 添加报价',
        'submit_text': '添加报价'
    }
    return render(request, 'sells/form.html', context)

def update_client_quotation(request, pk):
    quotation = get_object_or_404(SellsQuotation, pk=pk)
    if request.method == 'POST':
        form = SellsQuotationForm(request.POST, instance=quotation)
        if form.is_valid():
            form.save()
            messages.success(request, '设备更新成功')
            return redirect('client_detail', pk=quotation.client.pk)
    else:
        form = SellsQuotationForm(instance=quotation)
    context = {
        'form': form,
        'title': '更新报价',
        'submit_text': '更新报价'
    }
    return render(request, 'sells/form.html', context)

# ==================== 文档下载 ====================
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt # 用于设置字体大小
from docx.oxml.ns import qn # 用于设置中文字体
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

@login_required
#拜访记录下载
def generate_visit_record_docx(request, record_id):
    """
    根据给定的客户拜访记录ID生成并下载DOCX文档。
    文档实时生成，不保存到服务器。
    """
    visit_record = get_object_or_404(VisitRecord, pk=record_id)

    # 创建一个新的Word文档
    document = Document()

    # 定义一个辅助函数来添加带有指定字体的段落或标题
    def add_chinese_text(doc_obj, text, level=None, font_name='Microsoft YaHei', font_size=None):
        if level is not None:
            # 如果提供了level，使用add_heading来创建标题
            paragraph = doc_obj.add_heading(text, level=level)
        else:
            # 否则，使用add_paragraph来创建普通段落
            paragraph = doc_obj.add_paragraph(text)

        # 遍历段落中的所有run并设置字体
        # 有些情况下，add_heading或add_paragraph可能返回一个空的runs列表
        # 所以我们需要确保至少有一个run来设置字体
        if not paragraph.runs:
            run = paragraph.add_run()
            run.text = text # 重新设置文本，因为add_run()会清空它
        else:
            run = paragraph.runs[0] # 通常第一个run包含文本
            # 如果level是0，add_heading已经将文本放在了run里，不需要再次设置
            # 对于其他情况，可能需要确保文本被设置在run里
            # 为了确保一致性，我们可以选择在这里更新文本，或者依赖add_heading/add_paragraph的行为
            # 这里我们主要关注字体设置
            if level is None: # Only update text if it's a regular paragraph and not a heading
                 run.text = text


        font = run.font
        font.name = font_name
        # 设置中文字体
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            font.size = Pt(font_size)
        return paragraph

    # 添加标题和基本信息
    # 标题通常字号较大，可以使用不同的字体大小
    # 使用修改后的 add_chinese_text 函数
    add_chinese_text(document, '深圳市景曜数控设备有限公司', level=0, font_size=20) # level=0 for main title
    add_chinese_text(document, '客户拜访记录表', level=1, font_size=16) # level=1 for Heading 1

    add_chinese_text(document, '基本信息', level=2) # level=2 for Heading 2
    add_chinese_text(document, f"拜访日期: {visit_record.visit_date.strftime('%Y 年 %m 月 %d 日')}")
    add_chinese_text(document, f"拜访人员: {visit_record.salesperson.username}")
    add_chinese_text(document, f"客户名称: {visit_record.client.name}")
    add_chinese_text(document, f"客户联系人: {visit_record.client.contact_person}")
    add_chinese_text(document, f"联系电话: {visit_record.client.phone}")
    add_chinese_text(document, f"电子邮箱: {visit_record.client.email}")
    add_chinese_text(document, f"客户地址: {visit_record.client.address}")
    # add_chinese_text(document, f"所属行业: {visit_record.client.get_industry_display()}")

    # 一、拜访概况
    add_chinese_text(document, '一、拜访概况', level=2)
    add_chinese_text(document, f"拜访目的: {visit_record.get_purpose_display()}")
    add_chinese_text(document, f"拜访方式: {visit_record.get_method_display()}" + (f" ({visit_record.custom_method})" if visit_record.custom_method else ""))
    add_chinese_text(document, f"预计拜访时长: {visit_record.planned_duration} 小时")
    add_chinese_text(document, f"实际拜访时长: {visit_record.actual_duration} 小时")

    # 二、客户需求与反馈
    add_chinese_text(document, '二、客户需求与反馈', level=2)
    add_chinese_text(document, '客户现有设备情况', level=3) # level=3 for Heading 3
    # 获取客户设备信息
    client_equipments = ClientEquipment.objects.filter(client=visit_record.client)
    if client_equipments.exists():
        for i, equipment in enumerate(client_equipments):
            add_chinese_text(document, f"设备型号: {equipment.model}")
            add_chinese_text(document, f"使用年限: {equipment.years_in_use} 年")
            add_chinese_text(document, f"运行状况: {equipment.get_condition_display()} - {equipment.condition_description or '无描述'}")
            add_chinese_text(document, f"对现有设备的评价: {equipment.evaluation}")
            if i < len(client_equipments) - 1:
                add_chinese_text(document, "-" * 20) # 分隔符，用于多个设备
    else:
        add_chinese_text(document, "暂无设备信息")

    add_chinese_text(document, '客户新需求', level=3)
    add_chinese_text(document, f"是否有采购新设备计划: {'是' if visit_record.has_purchase_plan else '否'}")
    if visit_record.has_purchase_plan:
        add_chinese_text(document, f"计划采购时间: {visit_record.purchase_time.strftime('%Y-%m-%d') if visit_record.purchase_time else '未定'}")
        add_chinese_text(document, f"需求设备类型及规格: {visit_record.required_equipment or '无'}")
        add_chinese_text(document, f"特殊需求或关注点: {visit_record.special_requirements or '无'}")

    add_chinese_text(document, '客户对公司产品及服务的反馈', level=3)
    add_chinese_text(document, f"对公司现有设备产品的看法: {visit_record.feedback_on_products or '无'}")
    add_chinese_text(document, f"对售后服务的满意度: {visit_record.competitor_satisfaction or '未提及'}")
    add_chinese_text(document, f"其他意见或建议: {visit_record.other_comments or '无'}")

    # 三、销售进展
    add_chinese_text(document, '三、销售进展', level=2)
    add_chinese_text(document, '销售介绍内容', level=3)
    add_chinese_text(document, f"本次重点介绍的设备产品: {visit_record.featured_products}")
    add_chinese_text(document, f"产品优势与解决方案阐述: {visit_record.solutions_provided}")

    add_chinese_text(document, '客户合作意向', level=3)
    add_chinese_text(document, f"合作意向程度: {visit_record.get_cooperation_intention_display()}")
    add_chinese_text(document, f"下一步行动: {visit_record.next_steps}")

    # 四、竞争对手信息
    add_chinese_text(document, '四、竞争对手信息', level=2)
    client_competitors = Competitor.objects.filter(client=visit_record.client)
    if client_competitors.exists():
        add_chinese_text(document, "客户是否了解竞争对手产品: 是")
        add_chinese_text(document, '涉及竞争对手', level=3)
        for competitor in client_competitors:
            add_chinese_text(document, f"竞争对手名称: {competitor.name}")
            add_chinese_text(document, f"竞争对手产品优势: {competitor.advantages}")
            add_chinese_text(document, f"与本公司产品对比: {competitor.comparison}")
            add_chinese_text(document, "-" * 20)
    else:
        add_chinese_text(document, "客户是否了解竞争对手产品: 否")

    # 五、问题与建议
    add_chinese_text(document, '五、问题与建议', level=2)
    add_chinese_text(document, f"拜访过程中遇到的问题: {visit_record.issues_encountered or '无'}")
    add_chinese_text(document, f"针对问题的解决建议: {visit_record.proposed_solutions or '无'}")
    add_chinese_text(document, f"其他信息: {visit_record.other_message or '无'}")

    # 六、跟进计划
    add_chinese_text(document, '六、跟进计划', level=2)
    follow_up_plans = FollowUpPlan.objects.filter(visit_record=visit_record).order_by('follow_up_date')
    if follow_up_plans.exists():
        for i, plan in enumerate(follow_up_plans):
            add_chinese_text(document, f"跟进责任人: {plan.responsible_person.username}")
            add_chinese_text(document, f"第{i+1}次跟进时间: {plan.follow_up_date.strftime('%Y-%m-%d')}")
            add_chinese_text(document, f"预期达成目标: {plan.expected_outcome}")
            if i < len(follow_up_plans) - 1:
                add_chinese_text(document, "-" * 20)
    else:
        add_chinese_text(document, "暂无跟进计划")

    # 将文档保存到内存中的BytesIO对象
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0) # 将文件指针重置到开始

    # 创建HTTP响应，直接传输文件
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # 设置下载的文件名
    filename = f"客户拜访记录表_{visit_record.client.name}_{visit_record.visit_date.strftime('%Y%m%d')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response['Content-Length'] = length
    return response

#年度计划下载




@login_required
def download_annual_plan(request, pk):
    plan = get_object_or_404(AnnualPlan, pk=pk)

    document = Document()

    # --- 设置全局默认字体 (可选，但推荐) ---
    # 这会尝试为整个文档设置一个默认字体，特别是中文字体
    # 注意：这个方法可能不会对所有自动生成的段落生效，尤其是当它们有自己的样式时。
    # 更可靠的方法是在每个run上设置。
    # document.styles['Normal'].font.name = '宋体'
    # document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


    # --- 辅助函数：添加标题并设置字体和字号 ---
    def add_chinese_heading(doc, text, level, font_name='宋体', font_size_pt=16):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 设置中文字体
            run.font.size = Pt(font_size_pt) # 设置字号

    # --- 辅助函数：添加段落并设置字体和字号 ---
    def add_chinese_paragraph(doc, text, font_name='宋体', font_size_pt=10.5, alignment=None):
        paragraph = doc.add_paragraph(text)
        if alignment:
            paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 设置中文字体
            run.font.size = Pt(font_size_pt) # 设置字号

    # --- 设置公司名称和标题 ---
    add_chinese_heading(document, '深圳市景曜数控设备有限公司', level=1, font_size_pt=22) # 对应文档中的1级标题
    # 将标题居中
    for para in document.paragraphs:
        if para.text == '深圳市景曜数控设备有限公司':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    add_chinese_heading(document, f'{plan.year}年度销售计划表', level=2, font_size_pt=16) # 对应文档中的2级标题
    # 将标题居中
    for para in document.paragraphs:
        if para.text == f'{plan.year}年度销售计划表':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break


    # --- 一、计划周期 ---
    add_chinese_heading(document, '一、计划周期', level=3, font_size_pt=14)
    add_chinese_paragraph(document, f'{plan.year}年1月1日-{plan.year}年12月31日')

    # --- 二、销售目标 ---
    add_chinese_heading(document, '二、销售目标', level=3, font_size_pt=14)
    add_chinese_paragraph(document, f'总体目标 ：在本计划周期内，实现销售额达到 {plan.total_sales_target}万元.')
    add_chinese_paragraph(document, f'销售设备： SQZ系列钻靶机{plan.sqz_target}台；')
    add_chinese_paragraph(document, f'QZ系列钻靶机 {plan.qz_target}台；')
    add_chinese_paragraph(document, f'OL连线系列钻靶机 {plan.ol_target}台；')
    add_chinese_paragraph(document, f'Ro机器人系列钻靶机 {plan.ro_target}台；')
    add_chinese_paragraph(document, f'新客户开发数量达到 {plan.new_clients_target} 家，将内层板市场份额提升至 {plan.inner_layer_market_share}%，高端金属基板市场额提升到 {plan.high_end_metal_market_share}%。') # 假设你的模型里有这两个字段

    add_chinese_heading(document, '分解目标', level=4, font_size_pt=12)
    add_chinese_paragraph(document, '时间维度：将总销售额按季度分解，')
    add_chinese_paragraph(document, f'第一季度完成 {plan.q1_sales_target} 元，')
    add_chinese_paragraph(document, f'第二季度完成 {plan.q2_sales_target} 元，')
    add_chinese_paragraph(document, f'第三季度完成 {plan.q3_sales_target} 元，')
    add_chinese_paragraph(document, f'第四季度完成 {plan.q4_sales_target} 元；')
    add_chinese_paragraph(document, f'各季度销售设备数量分别为 {plan.q1_equipment_target} 台、{plan.q2_equipment_target} 台、{plan.q3_equipment_target} 台、{plan.q4_equipment_target} 台；')
    add_chinese_paragraph(document, f'新客户开发数量按季度依次为 {plan.q1_new_clients_target} 家、{plan.q2_new_clients_target} 家、{plan.q3_new_clients_target} 家、{plan.q4_new_clients_target} 家。')

    add_chinese_paragraph(document, '人员维度 ：销售经理负责监督整体销售目标达成，确保团队完成任务；销售主管根据团队成员能力与区域划分，将销售任务分配至每位销售人员。')
    # Assuming personnel_targets is a JSONField and you want to display it dynamically
    # 确保 personnel_targets 字段在你的 AnnualPlan 模型中存在且是 JSONField
    if plan.personnel_targets:
        for person, targets in plan.personnel_targets.items():
            add_chinese_paragraph(document, f'例如销售专员 {person} 需完成销售额 {targets.get("销售额", "N/A")} 元，销售设备 {targets.get("销售设备", "N/A")} 台，开发新客户 {targets.get("开发新客户", "N/A")} 家。')
    else:
        add_chinese_paragraph(document, '暂无人员维度目标数据。')


    # --- 三、市场分析 ---
    add_chinese_heading(document, '三、市场分析', level=3, font_size_pt=14)
    add_chinese_paragraph(document, f'行业动态 ：{plan.market_analysis}')
    add_chinese_paragraph(document, f'竞争对手分析： {plan.competitor_analysis}')
    add_chinese_paragraph(document, f'客户需求洞察： {plan.customer_insights}')

    # --- 四、销售策略 ---
    add_chinese_heading(document, '四、销售策略', level=3, font_size_pt=14)
    add_chinese_heading(document, '产品策略', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.product_strategy}')
    add_chinese_heading(document, '价格策略 ：', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.pricing_strategy}')
    add_chinese_heading(document, '渠道策略', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.channel_strategy}')
    add_chinese_heading(document, '促销策略', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.promotion_strategy}')

    # --- 五、执行计划 ---
    add_chinese_heading(document, '五、执行计划', level=3, font_size_pt=14)
    add_chinese_heading(document, '第一阶段（第 1 - 3 个月）', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.phase1_plan}')
    add_chinese_heading(document, '第二阶段（第 4 - 6 个月）', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.phase2_plan}')
    add_chinese_heading(document, '第三阶段（第 7 - 9 个月）', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.phase3_plan}')
    add_chinese_heading(document, '第四阶段（第 10 - 12 个月）', level=4, font_size_pt=12)
    add_chinese_paragraph(document, f'{plan.phase4_plan}')

    # --- 六、资源需求 ---
    add_chinese_heading(document, '六、资源需求', level=3, font_size_pt=14)
    add_chinese_paragraph(document, f'人力资源 ：{plan.hr_requirements}')
    add_chinese_paragraph(document, f'财务资源： {plan.financial_requirements}')
    add_chinese_paragraph(document, f'技术资源： {plan.technical_requirements}')

    # --- 七、风险评估与应对 ---
    add_chinese_heading(document, '七、风险评估与应对', level=3, font_size_pt=14)
    add_chinese_paragraph(document, f'市场风险 ：{plan.market_risks}')
    add_chinese_paragraph(document, f'竞争风险： {plan.competition_risks}')
    add_chinese_paragraph(document, f'客户信用风险： {plan.credit_risks}')
    add_chinese_paragraph(document, f'技术风险： {plan.technical_risks}')

    # --- 八、监督与评估 ---
    add_chinese_heading(document, '八、监督与评估', level=3, font_size_pt=14)
    add_chinese_paragraph(document, f'{plan.monitoring_plan}')

    # Save the document to a BytesIO object
    f = io.BytesIO()
    document.save(f)
    f.seek(0) # Rewind the buffer

    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{plan.year}年度销售计划.docx"'
    return response

# ... (之后的其他代码保持不变) ...
import csv
import re
#导入数据

def data_load(request):
    if request.method == 'GET':
        csv_file_path = 'app01/templates/sells/data/客户拜访记录表.csv'  # 修正文件扩展名为.csv

        print("开始导入数据...")
        try:
            with open(csv_file_path, 'r', encoding='utf-8-sig', errors='replace') as file:  # 添加errors参数处理编码问题
                reader = csv.DictReader(file)
                if not reader.fieldnames:  # 检查是否有表头
                    print("错误: CSV文件没有表头或为空")
                    return render(request, 'sells/data/load.html')

                for i, row in enumerate(reader):
                    # Skip rows where the company name is empty
                    company_name = row.get('公司名称', '').strip()
                    if not company_name:
                        print(f"Row {i + 2}: 跳过，缺少公司名称")
                        continue

                    try:
                        # 1. Create or get the Client object
                        client, client_created = Client.objects.get_or_create(
                            name=company_name,
                            defaults={'address': row.get('公司地址', '').strip()}
                        )
                        if not client_created and row.get('公司地址', '').strip():
                            client.address = row.get('公司地址', '').strip()
                            client.save()

                        # 2. Handle Contact 1 (updates the Client model directly)
                        contact1_text = row.get('联系人1', '').strip()
                        if contact1_text:
                            match = re.match(r'(.+?)(\d+)$', contact1_text)  # 添加$确保匹配到结尾
                            if match:
                                name, phone = match.groups()
                                client.contact_person = name.strip()
                                client.phone = phone.strip()
                            else:
                                client.contact_person = contact1_text
                                client.phone = ''
                            client.save()

                        # 3-4. Handle Contact 2 and 3 (saved in ClientContact)
                        for contact_field in ['联系人2', '联系人3']:
                            contact_text = row.get(contact_field, '').strip()
                            if contact_text:
                                match = re.match(r'(.+?)(\d+)$', contact_text)
                                if match:
                                    name, phone = match.groups()
                                    ClientContact.objects.get_or_create(
                                        client=client,
                                        name=name.strip(),
                                        defaults={'phone': phone.strip()}
                                    )
                                else:
                                    ClientContact.objects.get_or_create(
                                        client=client,
                                        name=contact_text
                                    )

                        # 5-8. Handle Products, Equipment, Purchase History and Quotations
                        model_field_map = {
                            '所做产品': ClientGeneration,
                            '现有设备': ClientEquipment,
                            '购买记录': ClientPurchase,
                            '报价': SellsQuotation
                        }

                        for field, model in model_field_map.items():
                            items = row.get(field, '').strip()
                            if items:
                                for item_name in filter(None, re.split(r'[、,，\s]+', items)):
                                    if field == '现有设备':
                                        model.objects.get_or_create(
                                            client=client,
                                            model=item_name.strip()
                                        )
                                    else:
                                        model.objects.get_or_create(
                                            client=client,
                                            name=item_name.strip()
                                        )

                        # 9. Combine remarks and visit log, then create VisitRecord
                        remarks = row.get('备注', '').strip()
                        visit_log_key = next((key for key in reader.fieldnames if '拜访记录' in key), None)
                        visit_log = row.get(visit_log_key, '').strip() if visit_log_key else ''

                        combined_message = ""
                        if remarks:
                            combined_message += f"备注：{remarks}"
                        if visit_log:
                            if combined_message:
                                combined_message += "\n"
                            combined_message += f"拜访记录：{visit_log}"

                        if combined_message:
                            VisitRecord.objects.create(
                                client=client,
                                visit_date=date.today(),
                                other_message=combined_message,
                                salesperson=request.user
                            )
                        print(f"Row {i + 2}: 成功处理客户 '{company_name}'")

                    except Exception as e:
                        print(f"Row {i + 2}: 处理客户 '{company_name}' 时出错 - {str(e)}")
                        continue

        except FileNotFoundError:
            print(f"错误: 文件 '{csv_file_path}' 未找到，请检查路径")
        except UnicodeDecodeError:
            print("错误: 文件编码问题，请确认文件使用GBK编码")
        except Exception as e:
            print(f"导入过程中发生错误: {e}")

    return render(request, 'sells/data/load.html')

